# # app/routers/candidature_rh.py
# from fastapi import APIRouter, HTTPException, Depends
# from sqlalchemy.orm import Session
# from app.db import get_db, engine
# from app.models.models import Candidature, Convocation
# from app.utils.pdf_generator import generate_convocation_pdf
# from datetime import datetime
# import sqlalchemy
# from sqlalchemy import text
# import os
# import logging
# import smtplib
# from email.mime.text import MIMEText
# from email.mime.multipart import MIMEMultipart
# from email.mime.base import MIMEBase
# from email import encoders
# import json
# import traceback
# from fastapi import Request


# router = APIRouter()
# logger = logging.getLogger(__name__)
# logging.basicConfig(level=logging.INFO)

# # ==========================================================
# # 🔹 GET candidatures (ORIGINAL + JOIN OFFRE + CONVOCATION)
# # ==========================================================
# @router.get("/candidatures")
# async def get_candidatures():
#     try:
#         query = sqlalchemy.text("""
#             SELECT 
#                 c.*,
#                 o.scoring_criteres,
#                 conv.date_entretien,
#                 conv.heure_entretien,
#                 conv.interval_minute,
#                 conv.lieu_entretien,
#                 c.telephone AS telephone_candidat,
#                 e.telephone AS telephone_employe
#             FROM candidatures c
#             LEFT JOIN offres o 
#                 ON o.id = c.offre_id
#             LEFT JOIN convocations conv
#                 ON conv.candidature_id = c.id
#             LEFT JOIN employes e 
#                 ON e.candidature_id = c.id
#             ORDER BY c.date_candidature DESC
#         """)

#         # ---------- Fonction nettoyage téléphone ----------
#         def clean_tel(tel):
#             if not tel:
#                 return None
#             if isinstance(tel, str) and tel.strip().lower() in [
#                 "non renseigné", "non renseigne", "n/a", "-", ""
#             ]:
#                 return None
#             return tel.strip()

#         # ---------- MinIO client setup ----------
#         from minio import Minio
#         MINIO_ENDPOINT = os.getenv("MINIO_ENDPOINT", "127.0.0.1:9000")
#         MINIO_ACCESS_KEY = os.getenv("MINIO_ROOT_USER", "jeremi")
#         MINIO_SECRET_KEY = os.getenv("MINIO_ROOT_PASSWORD", "Jeremi123")
#         MINIO_BUCKET = os.getenv("MINIO_BUCKET", "cvs")

#         minio_client = Minio(
#             MINIO_ENDPOINT,
#             access_key=MINIO_ACCESS_KEY,
#             secret_key=MINIO_SECRET_KEY,
#             secure=False
#         )

#         # ---------- Traitement ----------
#         import fitz  # PyMuPDF pour PDF
#         import docx  # python-docx pour DOCX

#         with engine.begin() as conn:
#             result = conn.execute(query)
#             candidatures = []
#             nouvelles_candidates = []

#             for row in result:
#                 r = dict(row._mapping)

#                 # ---------- NOM / PRENOM ----------
#                 if r.get("fullname"):
#                     parts = r["fullname"].split()
#                     r["nom"] = parts[0].upper()
#                     r["prenom"] = " ".join(parts[1:]) if len(parts) > 1 else ""

#                 # ---------- Candidature nouvelle ----------
#                 if not r.get("parsed_json"):
#                     nouvelles_candidates.append(r)

#                 # ---------- Téléphone ----------
#                 tel_candidat = clean_tel(r.get("telephone_candidat"))
#                 tel_employe = clean_tel(r.get("telephone_employe"))
#                 r["telephone"] = tel_candidat or tel_employe

#                 # ---------- Dates ----------
#                 r["date"] = r["date_candidature"].isoformat() if r.get("date_candidature") else None
#                 r["date_entretien"] = str(r.get("date_entretien")) if r.get("date_entretien") else None
#                 r["heure_entretien"] = str(r.get("heure_entretien")) if r.get("heure_entretien") else None

#                 candidatures.append(r)

#             # ---------- Parsing + scoring pour nouvelles candidatures ----------
#             for r in nouvelles_candidates:
#                 texte_cv = None

#                 # --- MinIO CV ---
#                 if r.get("raw_cv_s3"):
#                     try:
#                         object_name = r["raw_cv_s3"].split("/", 1)[1] if "/" in r["raw_cv_s3"] else r["raw_cv_s3"]
#                         response = minio_client.get_object(MINIO_BUCKET, object_name)
#                         file_bytes = response.read()
#                         from io import BytesIO
#                         bio = BytesIO(file_bytes)

#                         if object_name.lower().endswith(".pdf"):
#                             doc = fitz.open(stream=bio.read(), filetype="pdf")
#                             texte_cv = "\n".join([page.get_text() for page in doc])
#                         elif object_name.lower().endswith(".docx"):
#                             doc = docx.Document(bio)
#                             texte_cv = "\n".join([p.text for p in doc.paragraphs])
#                     except Exception:
#                         texte_cv = None

#                 # --- Fichier local ---
#                 elif r.get("cv_path") and os.path.exists(r["cv_path"]):
#                     try:
#                         path = r["cv_path"]
#                         if path.lower().endswith(".pdf"):
#                             doc = fitz.open(path)
#                             texte_cv = "\n".join([page.get_text() for page in doc])
#                         elif path.lower().endswith(".docx"):
#                             doc = docx.Document(path)
#                             texte_cv = "\n".join([p.text for p in doc.paragraphs])
#                     except Exception as e:
#                         print(f"❌ Erreur lecture CV local ID={r['id']}: {e}")
#                         texte_cv = None

#                 if not texte_cv:
#                     print(f"⚠️ Aucun texte CV trouvé pour ID={r['id']}")

#                 # --- Parse CV ---
#                 try:
#                     from app.utils.cv_parser import parse_cv_text
#                     parsed_cv = parse_cv_text(texte_cv) if texte_cv else {}
#                     if parsed_cv:
#                         update_parsed = sqlalchemy.text(
#                             "UPDATE candidatures SET parsed_json=:parsed WHERE id=:id"
#                         )
#                         conn.execute(update_parsed, {"parsed": json.dumps(parsed_cv, ensure_ascii=False), "id": r["id"]})
#                 except Exception:
#                     parsed_cv = {}

#                 # ---------- Score ----------
#                 offre_criteres = r.get("scoring_criteres") or {}
#                 score_total, breakdown = calculate_score(parsed_cv, offre_criteres)

#                 update_score = sqlalchemy.text(
#                     "UPDATE candidatures SET score_total=:score, score_breakdown=:breakdown WHERE id=:id"
#                 )
#                 conn.execute(update_score, {"score": score_total, "breakdown": json.dumps(breakdown, ensure_ascii=False), "id": r["id"]})

#                 r["parsed_json"] = parsed_cv
#                 r["score_total"] = score_total
#                 r["score_breakdown"] = breakdown
#                 r["score"] = score_total

#             # ---------- Sécurité pour anciennes candidatures ----------
#             for r in candidatures:
#                 if "score_total" not in r or r["score_total"] is None:
#                     r["score_total"] = r.get("score") or 0
#                 if "score_breakdown" not in r or not r["score_breakdown"]:
#                     r["score_breakdown"] = {}
#                 r["score"] = r["score_total"]

#         # ---------- Classement ----------
#         candidatures.sort(key=lambda x: x.get("score_total", 0), reverse=True)
#         print(f"📊 Total candidatures retournées: {len(candidatures)}")
#         return candidatures

#     except Exception as e:
#         print(traceback.format_exc())
#         raise HTTPException(status_code=500, detail=f"Erreur interne : {e}")

# # ==========================================================
# # 🔹 PUT sélection / désélection (ORIGINAL)
# # ==========================================================
# @router.put("/candidatures/{id}/select")
# async def select_candidature(id: int):
#     try:
#         query = sqlalchemy.text(
#             "UPDATE candidatures SET statut='Sélectionné' WHERE id=:id"
#         )
#         with engine.begin() as conn:
#             res = conn.execute(query, {"id": id})
#             if res.rowcount == 0:
#                 raise HTTPException(status_code=404, detail="Candidature non trouvée")
#         return {"message": "Candidature sélectionnée avec succès"}
#     except Exception as e:
#         print(traceback.format_exc())
#         raise HTTPException(status_code=500, detail=f"Erreur : {e}")

# @router.put("/candidatures/{id}/deselect")
# async def deselect_candidature(id: int):
#     try:
#         query = sqlalchemy.text(
#             "UPDATE candidatures SET statut='Désélectionné' WHERE id=:id"
#         )
#         with engine.begin() as conn:
#             res = conn.execute(query, {"id": id})
#             if res.rowcount == 0:
#                 raise HTTPException(status_code=404, detail="Candidature non trouvée")
#         return {"message": "Candidature désélectionnée avec succès"}
#     except Exception as e:
#         print(traceback.format_exc())
#         raise HTTPException(status_code=500, detail=f"Erreur : {e}")

# #
# # ==========================================================
# # 🔹 POST convocation + mail automatique (mandefa fotsiny rehefa sélectionné)
# # ==========================================================
# @router.post("/candidatures/{id}/send-invitation")
# async def send_invitation(id: int, db: Session = Depends(get_db)):
#     try:
#         # --- Vérification candidat ---
#         candidature = db.query(Candidature).filter(Candidature.id == id).first()
#         if not candidature:
#             raise HTTPException(status_code=404, detail="Candidature non trouvée")
#         if not candidature.email:
#             raise HTTPException(status_code=400, detail="Email du candidat manquant")
#         if candidature.statut != "Sélectionné":
#             raise HTTPException(status_code=400, detail="Seul les candidats sélectionnés peuvent recevoir une convocation")

#         # --- Création convocation si pas encore existante ---
#         now = datetime.now()
#         convocation = Convocation(
#             candidature_id=candidature.id,
#             date_entretien=now.strftime("%Y-%m-%d"),
#             heure_entretien=now.strftime("%H:%M"),
#             lieu_entretien="À définir",
#             status="en attente"
#         )
#         db.add(convocation)
#         db.commit()
#         db.refresh(convocation)

#         # --- Génération PDF ---
#         pdf_path = generate_convocation_pdf(candidature, convocation)
#         logger.info(f"✅ Convocation PDF généré : {pdf_path}")

#         # --- Récupération SMTP depuis DB ---
#         smtp_config = db.execute(text("SELECT * FROM smtp_config ORDER BY id DESC LIMIT 1")).fetchone()
#         if not smtp_config:
#             logger.error("❌ SMTP config non trouvée dans la DB")
#             raise HTTPException(status_code=500, detail="SMTP config tsy hita")

#         sender_email = smtp_config.email
#         sender_password = smtp_config.password
#         smtp_server = smtp_config.host or "smtp.gmail.com"
#         smtp_port = int(smtp_config.port or 587)
#         # --- Préparation email ---
#         message = MIMEMultipart()
#         message["Subject"] = "Convocation entretien - CODEL"
#         message["From"] = sender_email
#         message["To"] = candidature.email

#         html_content = f"""
#         <html>
#         <body>
#             <p>Bonjour <strong>{candidature.fullname}</strong>,</p>
#             <p>Vous êtes cordialement invité(e) à votre entretien pour le poste.</p>
#             <p><strong>Date et heure :</strong> {convocation.date_entretien} {convocation.heure_entretien}<br>
#                <strong>Lieu :</strong> {convocation.lieu_entretien}</p>
#             <p>Veuillez consulter le PDF joint pour tous les détails et apporter les documents nécessaires.</p>
#             <p><em>Message automatique, ne pas répondre à ce mail.</em></p>
#             <p>Cordialement,<br><strong>Équipe RH</strong></p>
#         </body>
#         </html>
#         """
#         message.attach(MIMEText(html_content, "html", "utf-8"))

#         with open(pdf_path, "rb") as f:
#             part = MIMEBase("application", "octet-stream")
#             part.set_payload(f.read())
#             encoders.encode_base64(part)
#             part.add_header(
#                 "Content-Disposition",
#                 f'attachment; filename="{os.path.basename(pdf_path)}"'
#             )
#             message.attach(part)

#         # --- Envoi email via SMTP DB ---
#         try:
#             with smtplib.SMTP(smtp_server, int(smtp_port)) as server:
#                 server.starttls()
#                 server.login(sender_email, sender_password)
#                 server.sendmail(sender_email, candidature.email, message.as_string())
#             logger.info(f"✅ Convocation envoyée à {candidature.fullname} ({candidature.email})")
#         except Exception as e:
#             logger.error(f"❌ Erreur SMTP : {str(e)}")
#             raise HTTPException(status_code=500, detail=f"Erreur lors de l'envoi du mail : {str(e)}")

#         # --- Mise à jour statut ---
#         convocation.status = "envoyée"
#         convocation.lien_fichier = pdf_path
#         candidature.statut = "Convoqué"
#         db.commit()

#         return {
#             "message": f"Bonjour {candidature.fullname}, votre convocation a été envoyée avec succès ✅",
#             "pdf_path": pdf_path,
#             "date_entretien": convocation.date_entretien,
#             "heure_entretien": convocation.heure_entretien,
#             "lieu_entretien": convocation.lieu_entretien,
#             "status": convocation.status
#         }

#     except Exception as e:
#         logger.error(traceback.format_exc())
#         raise HTTPException(status_code=500, detail=f"Erreur interne : {str(e)}")

# # ==========================================================
# # 🔹 CALCUL SCORE (SEUL CHANGEMENT)
# # ==========================================================
# def calculate_score(parsed_cv: dict, offre_criteres: dict) -> tuple[int, dict]:
#     """
#     ✔ Score basé sur critères OFFRE
#     ✔ Fallback automatique = logique originale
#     """

#     if not parsed_cv:
#         return 0, {}

#     scoring_config = offre_criteres if offre_criteres else {
#         "competences": ["Python", "SQL", "FastAPI"],
#         "experience_min": 3,
#         "diplome_requis": "Master Informatique",
#         "poids": {
#             "competences": 0.4,
#             "experience": 0.3,
#             "formation": 0.2,
#             "projets": 0.1
#         }
#     }

#     competences_cv = set(parsed_cv.get("competences", []))
#     competences_offre = set(scoring_config.get("competences", []))
#     match_comp = len(competences_cv & competences_offre) / max(1, len(competences_offre))

#     exp_cv = parsed_cv.get("experience_annees", 0)
#     exp_min = scoring_config.get("experience_min", 1)
#     match_exp = min(exp_cv / exp_min, 1)

#     formation_cv = parsed_cv.get("diplome", "")
#     match_formation = (
#         1 if formation_cv == scoring_config.get("diplome_requis") else 0.5
#     )

#     projets_cv = set(parsed_cv.get("projets", []))
#     match_proj = min(len(projets_cv), 3) / 3

#     poids = scoring_config.get("poids", {})
#     score_total = round(
#         (
#             match_comp * poids.get("competences", 0) +
#             match_exp * poids.get("experience", 0) +
#             match_formation * poids.get("formation", 0) +
#             match_proj * poids.get("projets", 0)
#         ) * 100
#     )

#     breakdown = {
#         "competences": round(match_comp * 100, 2),
#         "experience": round(match_exp * 100, 2),
#         "formation": round(match_formation * 100, 2),
#         "projets": round(match_proj * 100, 2),
#     }

#     return score_total, breakdown




# app/routers/candidature_rh.py
from fastapi import APIRouter, HTTPException, Depends
from sqlalchemy.orm import Session
from app.db import get_db, engine
from app.models.models import Candidature, Convocation
from app.utils.pdf_generator import generate_convocation_pdf
from datetime import datetime
import sqlalchemy
from sqlalchemy import text
import os
import logging
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import json
import traceback
from fastapi import Request


router = APIRouter()
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# ==========================================================
# 🔹 GET candidatures (ADAPTÉ pour vos tables)
# ==========================================================
@router.get("/candidatures")
async def get_candidatures():
    try:
        query = sqlalchemy.text("""
            SELECT 
                c.id,
                c.fullname,
                c.email,
                c.phone,
                c.adresse,
                c.date_naissance,
                c.poste,
                c.disponibilite,
                c.salaire,
                c.type_contrat,
                c.mobilite,
                c.autorisation,
                c.cv_path,
                c.lettre_path,
                c.diplomes_path,
                c.date_candidature,
                c.statut,
                c.score,
                c.offre_id,
                c.ref_offre,
                c.created_at,
                c.source,
                c.raw_cv_s3,
                c.parsed_json,
                json_build_object(
                    'scoring_config_path', o.scoring_config_path,
                    'tech_skills', COALESCE(o.tech_skills, '[]'::json),
                    'soft_skills', COALESCE(o.soft_skills, '[]'::json),
                    'education_level', o.education_level,
                    'exp_required_years', o.exp_required_years,
                    'w_skills', o.w_skills,
                    'w_exp', o.w_exp,
                    'w_edu', o.w_edu,
                    'w_proj', o.w_proj,
                    'threshold', o.threshold
                ) AS scoring_criteres,
                conv.date_entretien,
                conv.heure_entretien,
                conv.interval_minute,
                conv.lieu_entretien,
                c.phone AS telephone_candidat,
                e.telephone AS telephone_employe
            FROM candidatures c
            LEFT JOIN offres o 
                ON o.id = c.offre_id
            LEFT JOIN convocations conv
                ON conv.candidature_id = c.id
            LEFT JOIN employees e 
                ON e.candidature_id = c.id
            ORDER BY c.date_candidature DESC
        """)

        # ---------- Fonction nettoyage téléphone ----------
        def clean_tel(tel):
            if not tel:
                return None
            if isinstance(tel, str) and tel.strip().lower() in [
                "non renseigné", "non renseigne", "n/a", "-", ""
            ]:
                return None
            return tel.strip()

        # ---------- MinIO client setup ----------
        from minio import Minio
        MINIO_ENDPOINT = os.getenv("MINIO_ENDPOINT", "127.0.0.1:9000")
        MINIO_ACCESS_KEY = os.getenv("MINIO_ROOT_USER", "jeremi")
        MINIO_SECRET_KEY = os.getenv("MINIO_ROOT_PASSWORD", "Jeremi123")
        MINIO_BUCKET = os.getenv("MINIO_BUCKET", "cvs")

        minio_client = Minio(
            MINIO_ENDPOINT,
            access_key=MINIO_ACCESS_KEY,
            secret_key=MINIO_SECRET_KEY,
            secure=False
        )

        # ---------- Traitement ----------
        import fitz  # PyMuPDF pour PDF
        import docx  # python-docx pour DOCX

        with engine.begin() as conn:
            result = conn.execute(query)
            candidatures = []
            nouvelles_candidates = []

            for row in result:
                r = dict(row._mapping)

                # ---------- NOM / PRENOM ----------
                if r.get("fullname"):
                    parts = r["fullname"].split()
                    r["nom"] = parts[0].upper()
                    r["prenom"] = " ".join(parts[1:]) if len(parts) > 1 else ""

                # ---------- Candidature nouvelle ----------
                if not r.get("parsed_json"):
                    nouvelles_candidates.append(r)

                # ---------- Téléphone ----------
                tel_candidat = clean_tel(r.get("telephone_candidat"))
                tel_employe = clean_tel(r.get("telephone_employe"))
                r["telephone"] = tel_candidat or tel_employe

                # ---------- Dates ----------
                r["date"] = r["date_candidature"].isoformat() if r.get("date_candidature") else None
                r["date_entretien"] = str(r.get("date_entretien")) if r.get("date_entretien") else None
                r["heure_entretien"] = str(r.get("heure_entretien")) if r.get("heure_entretien") else None

                candidatures.append(r)

            # ---------- Parsing + scoring pour nouvelles candidatures ----------
            for r in nouvelles_candidates:
                texte_cv = None

                # --- MinIO CV ---
                if r.get("raw_cv_s3"):
                    try:
                        object_name = r["raw_cv_s3"].split("/", 1)[1] if "/" in r["raw_cv_s3"] else r["raw_cv_s3"]
                        response = minio_client.get_object(MINIO_BUCKET, object_name)
                        file_bytes = response.read()
                        from io import BytesIO
                        bio = BytesIO(file_bytes)

                        if object_name.lower().endswith(".pdf"):
                            doc = fitz.open(stream=bio.read(), filetype="pdf")
                            texte_cv = "\n".join([page.get_text() for page in doc])
                        elif object_name.lower().endswith(".docx"):
                            doc = docx.Document(bio)
                            texte_cv = "\n".join([p.text for p in doc.paragraphs])
                    except Exception:
                        texte_cv = None

                # --- Fichier local ---
                elif r.get("cv_path") and os.path.exists(r["cv_path"]):
                    try:
                        path = r["cv_path"]
                        if path.lower().endswith(".pdf"):
                            doc = fitz.open(path)
                            texte_cv = "\n".join([page.get_text() for page in doc])
                        elif path.lower().endswith(".docx"):
                            doc = docx.Document(path)
                            texte_cv = "\n".join([p.text for p in doc.paragraphs])
                    except Exception as e:
                        print(f"❌ Erreur lecture CV local ID={r['id']}: {e}")
                        texte_cv = None

                if not texte_cv:
                    print(f"⚠️ Aucun texte CV trouvé pour ID={r['id']}")

                # --- Parse CV ---
                try:
                    from app.utils.cv_parser import parse_cv_text
                    parsed_cv = parse_cv_text(texte_cv) if texte_cv else {}
                    if parsed_cv:
                        update_parsed = sqlalchemy.text(
                            "UPDATE candidatures SET parsed_json=:parsed WHERE id=:id"
                        )
                        conn.execute(update_parsed, {"parsed": json.dumps(parsed_cv, ensure_ascii=False), "id": r["id"]})
                except Exception:
                    parsed_cv = {}

                # ---------- Score ----------
                offre_criteres = r.get("scoring_criteres") or {}
                score_total, breakdown = calculate_score(parsed_cv, offre_criteres)

                # Mise à jour dans la base
                update_score = sqlalchemy.text(
                    "UPDATE candidatures SET score=:score WHERE id=:id"
                )
                conn.execute(update_score, {
                    "score": score_total, 
                    "id": r["id"]
                })

                r["parsed_json"] = parsed_cv
                r["score_total"] = score_total
                r["score_breakdown"] = breakdown
                r["score"] = score_total

            # ---------- Sécurité pour anciennes candidatures ----------
            for r in candidatures:
                if "score_total" not in r or r["score_total"] is None:
                    r["score_total"] = r.get("score") or 0
                if "score_breakdown" not in r or not r["score_breakdown"]:
                    r["score_breakdown"] = {}
                r["score"] = r["score_total"]

        # ---------- Classement ----------
        candidatures.sort(key=lambda x: x.get("score_total", 0), reverse=True)
        print(f"📊 Total candidatures retournées: {len(candidatures)}")
        return candidatures

    except Exception as e:
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Erreur interne : {e}")

# ==========================================================
# 🔹 PUT sélection / désélection
# ==========================================================
@router.put("/candidatures/{id}/select")
async def select_candidature(id: int):
    try:
        query = sqlalchemy.text(
            "UPDATE candidatures SET statut='Sélectionné' WHERE id=:id"
        )
        with engine.begin() as conn:
            res = conn.execute(query, {"id": id})
            if res.rowcount == 0:
                raise HTTPException(status_code=404, detail="Candidature non trouvée")
        return {"message": "Candidature sélectionnée avec succès"}
    except Exception as e:
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Erreur : {e}")

@router.put("/candidatures/{id}/deselect")
async def deselect_candidature(id: int):
    try:
        query = sqlalchemy.text(
            "UPDATE candidatures SET statut='Désélectionné' WHERE id=:id"
        )
        with engine.begin() as conn:
            res = conn.execute(query, {"id": id})
            if res.rowcount == 0:
                raise HTTPException(status_code=404, detail="Candidature non trouvée")
        return {"message": "Candidature désélectionnée avec succès"}
    except Exception as e:
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Erreur : {e}")

# ==========================================================
# 🔹 POST convocation + mail automatique
# ==========================================================
@router.post("/candidatures/{id}/send-invitation")
async def send_invitation(id: int, db: Session = Depends(get_db)):
    try:
        # --- Vérification candidat ---
        candidature = db.query(Candidature).filter(Candidature.id == id).first()
        if not candidature:
            raise HTTPException(status_code=404, detail="Candidature non trouvée")
        if not candidature.email:
            raise HTTPException(status_code=400, detail="Email du candidat manquant")
        if candidature.statut != "Sélectionné":
            raise HTTPException(status_code=400, detail="Seul les candidats sélectionnés peuvent recevoir une convocation")

        # --- Création convocation si pas encore existante ---
        now = datetime.now()
        convocation = Convocation(
            candidature_id=candidature.id,
            date_entretien=now.strftime("%Y-%m-%d"),
            heure_entretien=now.strftime("%H:%M"),
            lieu_entretien="À définir",
            status="en attente"
        )
        db.add(convocation)
        db.commit()
        db.refresh(convocation)

        # --- Génération PDF ---
        pdf_path = generate_convocation_pdf(candidature, convocation)
        logger.info(f"✅ Convocation PDF généré : {pdf_path}")

        # --- Récupération SMTP depuis DB ---
        smtp_config = db.execute(text("SELECT * FROM smtp_config ORDER BY id DESC LIMIT 1")).fetchone()
        if not smtp_config:
            logger.error("❌ SMTP config non trouvée dans la DB")
            raise HTTPException(status_code=500, detail="SMTP config tsy hita")

        sender_email = smtp_config.email
        sender_password = smtp_config.password
        smtp_server = smtp_config.host or "smtp.gmail.com"
        smtp_port = int(smtp_config.port or 587)
        # --- Préparation email ---
        message = MIMEMultipart()
        message["Subject"] = "Convocation entretien - CODEL"
        message["From"] = sender_email
        message["To"] = candidature.email

        html_content = f"""
        <html>
        <body>
            <p>Bonjour <strong>{candidature.fullname}</strong>,</p>
            <p>Vous êtes cordialement invité(e) à votre entretien pour le poste.</p>
            <p><strong>Date et heure :</strong> {convocation.date_entretien} {convocation.heure_entretien}<br>
               <strong>Lieu :</strong> {convocation.lieu_entretien}</p>
            <p>Veuillez consulter le PDF joint pour tous les détails et apporter les documents nécessaires.</p>
            <p><em>Message automatique, ne pas répondre à ce mail.</em></p>
            <p>Cordialement,<br><strong>Équipe RH</strong></p>
        </body>
        </html>
        """
        message.attach(MIMEText(html_content, "html", "utf-8"))

        with open(pdf_path, "rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                f'attachment; filename="{os.path.basename(pdf_path)}"'
            )
            message.attach(part)

        # --- Envoi email via SMTP DB ---
        try:
            with smtplib.SMTP(smtp_server, int(smtp_port)) as server:
                server.starttls()
                server.login(sender_email, sender_password)
                server.sendmail(sender_email, candidature.email, message.as_string())
            logger.info(f"✅ Convocation envoyée à {candidature.fullname} ({candidature.email})")
        except Exception as e:
            logger.error(f"❌ Erreur SMTP : {str(e)}")
            raise HTTPException(status_code=500, detail=f"Erreur lors de l'envoi du mail : {str(e)}")

        # --- Mise à jour statut ---
        convocation.status = "envoyée"
        convocation.lien_fichier = pdf_path
        candidature.statut = "Convoqué"
        db.commit()

        return {
            "message": f"Bonjour {candidature.fullname}, votre convocation a été envoyée avec succès ✅",
            "pdf_path": pdf_path,
            "date_entretien": convocation.date_entretien,
            "heure_entretien": convocation.heure_entretien,
            "lieu_entretien": convocation.lieu_entretien,
            "status": convocation.status
        }

    except Exception as e:
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Erreur interne : {str(e)}")

# ==========================================================
# 🔹 CALCUL SCORE (ADAPTÉ pour structure offres)
# ==========================================================
def calculate_score(parsed_cv: dict, offre_criteres: dict) -> tuple[int, dict]:
    """
    ✔ Score basé sur critères OFFRE (nouvelle structure)
    """
    if not parsed_cv:
        return 0, {}

    # Raha offre_criteres dia JSON string
    if isinstance(offre_criteres, str):
        try:
            offre_criteres = json.loads(offre_criteres) if offre_criteres else {}
        except:
            offre_criteres = {}

    scoring_config = offre_criteres if offre_criteres else {
        "tech_skills": ["Python", "SQL", "FastAPI"],
        "exp_required_years": 3,
        "education_level": "Master Informatique",
        "w_skills": 0.4,
        "w_exp": 0.3,
        "w_edu": 0.2,
        "w_proj": 0.1
    }

    # Compétences techniques
    competences_cv = set(parsed_cv.get("competences", []))
    tech_skills_offre = set(scoring_config.get("tech_skills", []))
    match_comp = len(competences_cv & tech_skills_offre) / max(1, len(tech_skills_offre))

    # Expérience
    exp_cv = parsed_cv.get("experience_annees", 0)
    exp_min = scoring_config.get("exp_required_years", 1)
    match_exp = min(exp_cv / exp_min, 1) if exp_min > 0 else 0

    # Formation
    formation_cv = parsed_cv.get("diplome", "")
    formation_requise = scoring_config.get("education_level", "")
    match_formation = (
        1 if formation_cv and formation_requise and formation_requise.lower() in formation_cv.lower() else 0.5
    )

    # Projets
    projets_cv = set(parsed_cv.get("projets", []))
    match_proj = min(len(projets_cv), 3) / 3

    # Poids
    w_skills = scoring_config.get("w_skills", 0.4)
    w_exp = scoring_config.get("w_exp", 0.3)
    w_edu = scoring_config.get("w_edu", 0.2)
    w_proj = scoring_config.get("w_proj", 0.1)

    score_total = round(
        (
            match_comp * w_skills +
            match_exp * w_exp +
            match_formation * w_edu +
            match_proj * w_proj
        ) * 100
    )

    breakdown = {
        "competences": round(match_comp * 100, 2),
        "experience": round(match_exp * 100, 2),
        "formation": round(match_formation * 100, 2),
        "projets": round(match_proj * 100, 2),
    }

    return score_total, breakdown
