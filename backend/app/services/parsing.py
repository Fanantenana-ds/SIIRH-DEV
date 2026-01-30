# from docx import Document
# import pdfplumber
# import re

# # --- 1️⃣ Parser DOCX ---
# def parse_docx(path: str) -> str:
#     doc = Document(path)
#     text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
#     return text

# # --- 2️⃣ Parser PDF ---
# def parse_pdf(path: str) -> str:
#     text = ""
#     with pdfplumber.open(path) as pdf:
#         for page in pdf.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"
#     return text

# # --- 3️⃣ Extraction automatique d'informations ---
# def extract_info(text: str, project_keywords: list = None) -> dict:
#     """
#     Extrait les informations principales du CV.
#     Optionnel: project_keywords pour extraire mots-clés projets spécifiques.
#     """
#     # Email
#     email = re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b", text)

#     # Téléphone
#     phone = re.search(r"\+?\d[\d\s().-]{7,}", text)

#     # Nom/Prénom
#     name_match = re.findall(r"([A-Z][a-z]+)", text)
#     firstname, lastname = (name_match[0], name_match[1]) if len(name_match) >= 2 else (None, None)

#     # Compétences générales
#     skill_keywords = [
#         "Python", "Django", "FastAPI", "React", "SQL", "Docker", "Kubernetes",
#         "Machine Learning", "Data Analysis", "Excel", "Communication", "Leadership"
#     ]
#     skills = [kw for kw in skill_keywords if re.search(rf"\b{kw}\b", text, re.IGNORECASE)]

#     # Diplômes
#     diploma_keywords = ["Licence", "Master", "Doctorat", "Ingénieur", "Bachelor"]
#     diplomes = [d for d in diploma_keywords if re.search(rf"\b{d}\b", text, re.IGNORECASE)]

#     # Langues
#     langues_keywords = ["Français", "Anglais", "Espagnol", "Allemand", "Italien", "Malgache"]
#     langues = [l for l in langues_keywords if re.search(rf"\b{l}\b", text, re.IGNORECASE)]

#     # Expérience
#     exp_years_match = re.findall(r"(\d+)\s+(ans|années|an)", text.lower())
#     exp_years = max([int(e[0]) for e in exp_years_match], default=0)

#     # --- Extraction mots-clés projets spécifiques ---
#     project_matches = []
#     if project_keywords:
#         for kw in project_keywords:
#             if re.search(rf"\b{re.escape(kw)}\b", text, re.IGNORECASE):
#                 project_matches.append(kw)

#     return {
#         "firstname": firstname,
#         "lastname": lastname,
#         "email": email.group() if email else None,
#         "phone": phone.group() if phone else None,
#         "skills": skills,
#         "diplomes": diplomes,
#         "langues": langues,
#         "exp_years": exp_years,
#         "projects": project_matches,  # nouveauté pour scoring automatique
#         "text": text[:3000],  # résumé pour scoring
#     }

# # --- 4️⃣ Extraction nom + téléphone simple ---
# def extract_name_phone(text: str) -> dict:
#     """
#     Fonction dédiée pour extraire le nom complet et le téléphone.
#     Nécessaire pour upload_service.py afin d'éviter ImportError.
#     """
#     name = None
#     phone = None

#     # Exemple rapide: prend les 2 premiers mots avec majuscule pour nom/prénom
#     name_match = re.findall(r"[A-Z][a-z]+", text)
#     if len(name_match) >= 2:
#         name = f"{name_match[0]} {name_match[1]}"

#     # Téléphone simple
#     phone_match = re.search(r"\+?\d[\d\s().-]{7,}", text)
#     if phone_match:
#         phone = phone_match.group()

#     return {
#         "fullname": name,
#         "phone": phone
#     }





# app/services/parsing.py - VERSION MONDE RÉEL COMPLÈTE
import io
import re
import logging
import json
from typing import Dict, Any, List, Optional, Tuple
import traceback
from datetime import datetime

logger = logging.getLogger(__name__)

# ================================
# 📚 IMPORTS AVEC FALLBACK ROBUSTE
# ================================
def safe_import(module_name: str, install_cmd: str = ""):
    """Import sécurisé avec gestion d'erreurs"""
    try:
        if module_name == "pdfplumber":
            import pdfplumber
            return pdfplumber, True
        elif module_name == "spacy":
            import spacy
            return spacy, True
        elif module_name == "langdetect":
            from langdetect import detect
            return detect, True
        elif module_name == "docx":
            from docx import Document
            return Document, True
        elif module_name == "pytesseract":
            import pytesseract
            from PIL import Image
            return (pytesseract, Image), True
        elif module_name == "PyPDF2":
            import PyPDF2
            return PyPDF2, True
    except ImportError as e:
        logger.warning(f"⚠️ {module_name} non disponible: {install_cmd}")
        return None, False

# Imports sécurisés
pdfplumber_module, PDFPLUMBER_AVAILABLE = safe_import("pdfplumber", "pip install pdfplumber")
spacy_module, SPACY_AVAILABLE = safe_import("spacy", "pip install spacy")
detect_func, LANGDETECT_AVAILABLE = safe_import("langdetect", "pip install langdetect")
Document_class, DOCX_AVAILABLE = safe_import("docx", "pip install python-docx")
ocr_modules, OCR_AVAILABLE = safe_import("pytesseract", "pip install pytesseract pillow")
pypdf_module, PDF2_AVAILABLE = safe_import("PyPDF2", "pip install PyPDF2")

# ================================
# 🔥 CHARGEMENT MODÈLES SPAcy
# ================================
nlp = None
if SPACY_AVAILABLE:
    try:
        nlp = spacy_module.load("en_core_web_sm")
        logger.info("✅ spaCy model 'en_core_web_sm' chargé")
    except OSError:
        try:
            nlp = spacy_module.load("fr_core_news_sm")
            logger.info("✅ spaCy model 'fr_core_news_sm' chargé")
        except OSError as e:
            logger.warning(f"⚠️ Aucun modèle spaCy: {e}")
            nlp = None
else:
    logger.info("ℹ️ spaCy non disponible, utilisation extraction regex")

# ================================
# 🧠 LISTES DE MOTS-CLÉS POUR CV MALGACHE/FRANÇAIS
# ================================
MALAGASY_NAMES = [
    'rajaona', 'randrianarisoa', 'rakotomalala', 'randriamanantsoa', 'razafindrakoto',
    'andriamalala', 'rasoanaivo', 'randrianarivony', 'randriamiarintsoa', 'raharison',
    'ramaroson', 'randrianarisoa', 'rakotondrabe', 'razafindrabe', 'randrianantoandro',
    'rakoto', 'rasoa', 'raoelina', 'randria', 'razafy', 'andry', 'hery', 'niry', 'tiana',
    'faniry', 'harivelo', 'tahina', 'lova', 'hasina', 'nirina', 'voahangy', 'sarobidy'
]

FRENCH_NAMES = [
    'martin', 'bernard', 'dubois', 'thomas', 'robert', 'richard', 'petit', 'durand',
    'leroy', 'moreau', 'simon', 'laurent', 'lefebvre', 'michel', 'garcia', 'david',
    'bertrand', 'roux', 'vincent', 'fournier', 'morel', 'girard', 'andre', 'lefevre',
    'mercier', 'dupont', 'lambert', 'bonnet', 'francois', 'martinez',
    'jean', 'pierre', 'paul', 'jacques', 'michel', 'alain', 'patrick', 'nicolas',
    'christophe', 'daniel', 'rene', 'eric', 'stephane', 'david', 'frederic',
    'marie', 'anne', 'isabelle', 'sophie', 'catherine', 'francoise', 'monique',
    'nathalie', 'christine', 'valerie', 'sandrine', 'caroline', 'patricia'
]

# ================================
# 🏙️ LISTE ANARANA TANÀNA MALAGASY
# ================================
MALAGASY_TOWN_NAMES = [
    'fianarantsoa', 'toliara', 'toamasina', 'mahajanga', 'antsiranana',
    'antananarivo', 'morondava', 'taolagnaro', 'antsirabe', 'ambositra',
    'moramanga', 'tamatave', 'majunga', 'diego', 'tulear', 'fort dauphin',
    'manakara', 'sambava', 'maroantsetra', 'miandrivazo', 'ambatondrazaka',
    'morombe', 'bekily', 'ambovombe', 'ihosy', 'farafangana', 'mananjary',
    'mahanoro', 'vatomandry', 'marolambo', 'mananara', 'andapa', 'vohemar',
    'soanierana', 'ivato', 'ambohidratrimo', 'alarobia', 'ambohimanga',
    'analakely', 'andoharanofotsy', 'ankadikely', 'ankadindramamy',
    'ankaditapaka', 'ankadivato', 'ankorondrano', 'ankotika', 'antsahabe',
    'antsakaviro', 'behoririka', 'faravohitra', 'isotry', 'mahamasina',
    'manjakandriana', 'masindray', 'namehana', 'soavinandriana', 'tambohobe',
    'tsaralalana', 'ampahitrosy', 'ampasampito', 'andranomena', 'ankadimanga',
    'ankadinandriana', 'ankaditany', 'ankoronana', 'antsararay', 'faribolana',
    'mandroseza', 'merimanjaka', 'soamanandray', 'talatamaty', 'tsararay',
    'ambatolampy', 'ambatomainty', 'ambatomena', 'ambodifototra', 'ambodivoanio',
    'ampanotokana', 'andilamena', 'andranofasika', 'ankazobe', 'ankorabe',
    'antsalova', 'antsampandrano', 'betafo', 'fenoarivo', 'mahabo', 'manandona',
    'mananivo', 'mandabe', 'mandritsara', 'manja', 'maroala', 'marovoay',
    'miatrika', 'mitsinjo', 'nosy be', 'sahatona', 'sakaraha', 'soalala',
    'tolanaro', 'tsiombe', 'vangaindrano', 'vohibinany', 'vohimarina', 'zoma'
]

TECH_SKILLS = {
    'programming': ['python', 'java', 'javascript', 'php', 'c#', 'c++', 'ruby', 'go', 'swift', 'kotlin'],
    'web': ['html', 'css', 'react', 'angular', 'vue', 'node.js', 'django', 'flask', 'spring', 'laravel'],
    'database': ['sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'oracle', 'sqlite', 'mariadb'],
    'data_science': ['pandas', 'numpy', 'scikit-learn', 'tensorflow', 'pytorch', 'keras', 'matplotlib', 'seaborn'],
    'devops': ['docker', 'kubernetes', 'aws', 'azure', 'gcp', 'jenkins', 'git', 'ci/cd', 'ansible', 'terraform'],
    'business': ['power bi', 'tableau', 'excel', 'vba', 'sap', 'oracle', 'erp', 'crm'],
    'soft_skills': ['communication', 'leadership', 'travail équipe', 'autonomie', 'rigueur', 'créativité']
}

# ================================
# 🧩 EXTRACTION TEXTE AMÉLIORÉE
# ================================
def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """Extraction texte robuste pour différents formats"""
    filename = filename.lower()
    
    # PDF
    if filename.endswith(".pdf"):
        # Essayer pdfplumber d'abord
        if PDFPLUMBER_AVAILABLE:
            try:
                text = _extract_with_pdfplumber(content)
                if text and len(text.strip()) > 100:
                    return text
            except Exception as e:
                logger.warning(f"pdfplumber échoué: {e}")
        
        # Fallback: PyPDF2
        if PDF2_AVAILABLE:
            try:
                text = _extract_with_pypdf2(content)
                if text and len(text.strip()) > 100:
                    return text
            except Exception as e:
                logger.warning(f"PyPDF2 échoué: {e}")
        
        # Dernier recours: OCR si disponible
        if OCR_AVAILABLE and len(content) < 10 * 1024 * 1024:  # Max 10MB pour OCR
            try:
                text = _extract_with_ocr(content)
                if text:
                    return text
            except Exception as e:
                logger.warning(f"OCR échoué: {e}")
    
    # DOCX
    elif filename.endswith(".docx") and DOCX_AVAILABLE:
        try:
            return _extract_docx(content)
        except Exception as e:
            logger.warning(f"DOCX échoué: {e}")
    
    # TXT et autres
    return _extract_text_fallback(content)

def _extract_with_pdfplumber(content: bytes) -> str:
    """Extraction avec pdfplumber"""
    text = ""
    with pdfplumber_module.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages):
            if i >= 10:  # Limiter à 10 pages
                break
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def _extract_with_pypdf2(content: bytes) -> str:
    """Extraction avec PyPDF2 (fallback)"""
    text = ""
    pdf_reader = pypdf_module.PdfReader(io.BytesIO(content))
    for i, page in enumerate(pdf_reader.pages):
        if i >= 10:
            break
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def _extract_with_ocr(content: bytes) -> str:
    """OCR pour PDF scannés"""
    text = ""
    pytesseract, Image = ocr_modules
    try:
        # Convertir PDF en images (première page seulement)
        if PDFPLUMBER_AVAILABLE:
            with pdfplumber_module.open(io.BytesIO(content)) as pdf:
                if pdf.pages:
                    page = pdf.pages[0]
                    img = page.to_image(resolution=150).original
                    text = pytesseract.image_to_string(img, lang='fra+eng')
    except Exception as e:
        logger.error(f"OCR error: {e}")
    return text

def _extract_docx(content: bytes) -> str:
    """Extraction DOCX"""
    doc = Document_class(io.BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def _extract_text_fallback(content: bytes) -> str:
    """Fallback simple"""
    for encoding in ['utf-8', 'latin-1', 'iso-8859-1', 'windows-1252']:
        try:
            return content.decode(encoding, errors='ignore')
        except:
            continue
    return ""

# ================================
# 🔍 EXTRACTION NOM COMPLET AVEC FANAVAHANA TANÀNA
# ================================
def extract_fullname(text: str) -> str:
    """Extraction avancée du nom avec filtrage des noms de villes"""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    # ====================
    # ÉTAPE 1: FILTRAGE TANÀNA
    # ====================
    filtered_lines = []
    for line in lines[:15]:  # Seulement les 15 premières lignes
        line_lower = line.lower()
        
        # Ignorer si c'est un nom de ville
        is_town_name = False
        for town in MALAGASY_TOWN_NAMES:
            if town in line_lower:
                is_town_name = True
                logger.debug(f"   ⛔ Ignorer ligne (tanàna): {line}")
                break
        
        # Ignorer les lignes avec mots-clés d'adresse
        address_keywords = ['rue', 'avenue', 'boulevard', 'lot', 'immeuble', 
                          'bp', 'boite postale', 'postale', 'adresse', 'quartier',
                          'district', 'province', 'region', 'ville', 'tanàna',
                          'commune', 'fokontany', 'madagascar', 'mg', 'address',
                          'street', 'avenue', 'boulevard', 'postal', 'code']
        
        has_address = any(keyword in line_lower for keyword in address_keywords)
        
        # Ignorer les titres seuls
        is_title_only = line_lower in ['mr', 'mme', 'monsieur', 'madame', 'mademoiselle']
        
        if not is_town_name and not has_address and not is_title_only and len(line) >= 3:
            filtered_lines.append(line)
    
    # ====================
    # ÉTAPE 2: RECHERCHE PATTERN NOM OLO
    # ====================
    for line in filtered_lines[:10]:  # 10 premières lignes filtrées
        line_clean = re.sub(r'[^\w\s\-\']', '', line)
        words = [w.strip() for w in line_clean.split() if w.strip()]
        
        if len(words) >= 2 and len(words) <= 4:
            # Pattern 1: "RAKOTO Jean" (NOM majuscule, prénom Capitalize)
            if (words[0].isupper() and len(words[0]) > 2 and 
                words[1][0].isupper() and words[1][1:].islower()):
                logger.debug(f"   ✅ Pattern 1 trouvé: {' '.join(words)}")
                return ' '.join(words)
            
            # Pattern 2: "Jean RAKOTO" (prénom puis NOM)
            if (words[-1].isupper() and len(words[-1]) > 2 and
                words[0][0].isupper() and words[0][1:].islower()):
                logger.debug(f"   ✅ Pattern 2 trouvé: {' '.join(words)}")
                return ' '.join(words)
            
            # Pattern 3: Deux mots, tous deux Capitalize (sans majuscule complète)
            if (len(words) == 2 and 
                words[0][0].isupper() and words[0][1:].islower() and
                words[1][0].isupper() and words[1][1:].islower()):
                # Vérifier que ce ne sont pas des villes connues
                if (words[0].lower() not in MALAGASY_TOWN_NAMES and 
                    words[1].lower() not in MALAGASY_TOWN_NAMES):
                    logger.debug(f"   ✅ Pattern 3 trouvé: {' '.join(words)}")
                    return ' '.join(words)
    
    # ====================
    # ÉTAPE 3: RECHERCHE SECTION SPÉCIFIQUE
    # ====================
    section_patterns = [
        r'(?i)^\s*(?:nom[\s:]*|name[\s:]*)\s*(.+)$',
        r'(?i)^\s*(?:candidat[\s:]*|applicant[\s:]*)\s*(.+)$',
        r'(?i)^\s*(?:profil[\s:]*|profile[\s:]*)\s*(.+)$',
        r'(?i)^\s*(.+?)\s*$[\n]*^\s*(?:contact|coordonnées|informations personnelles)\s*$',
    ]
    
    for pattern in section_patterns:
        matches = re.findall(pattern, text[:1500], re.MULTILINE)
        for match in matches:
            candidate = match.strip()
            candidate_lower = candidate.lower()
            
            # Vérifier que ce n'est pas une ville
            is_town = any(town in candidate_lower for town in MALAGASY_TOWN_NAMES[:20])
            has_address_word = any(word in candidate_lower for word in ['rue', 'avenue', 'lot'])
            
            if not is_town and not has_address_word and len(candidate.split()) >= 2:
                logger.debug(f"   ✅ Section trouvée: {candidate}")
                return candidate
    
    # ====================
    # ÉTAPE 4: SPAcy NER AVEC FILTRE
    # ====================
    if nlp and len(text) > 100:
        try:
            doc = nlp(text[:1500])
            persons = []
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    ent_lower = ent.text.lower()
                    # Vérifier que ce n'est pas une ville
                    is_town = False
                    for town in MALAGASY_TOWN_NAMES:
                        if town in ent_lower:
                            is_town = True
                            break
                    
                    if not is_town:
                        persons.append(ent.text)
            
            if persons:
                logger.debug(f"   ✅ SPAcy NER trouvé: {persons[0]}")
                return persons[0]
        except Exception as e:
            logger.debug(f"spaCy NER échoué: {e}")
    
    # ====================
    # ÉTAPE 5: RECHERCHE PAR PRÉNOMS CONNUS
    # ====================
    text_lower = text.lower()
    all_names = MALAGASY_NAMES + FRENCH_NAMES
    
    for name in all_names:
        if name in text_lower:
            # Trouver le contexte autour du prénom
            pattern = r'[\n\.]?[^\n\.]*\b' + re.escape(name) + r'\b[^\n\.]*'
            matches = re.findall(pattern, text, re.IGNORECASE)
            
            for match in matches:
                match_strip = match.strip()
                if len(match_strip.split()) >= 2:
                    # Vérifier que ce n'est pas une ville
                    match_lower = match_strip.lower()
                    is_town = any(town in match_lower for town in MALAGASY_TOWN_NAMES)
                    
                    if not is_town:
                        logger.debug(f"   ✅ Prénom connu trouvé: {match_strip}")
                        return match_strip
    
    # ====================
    # ÉTAPE 6: EXTRACTION EMAIL
    # ====================
    email_match = re.search(r'([a-zA-Z0-9._%+-]+)@', text)
    if email_match:
        email_part = email_match.group(1)
        # Convertir "jean.rakoto" en "Jean Rakoto"
        name_parts = re.sub(r'[._0-9]+', ' ', email_part).split()
        if len(name_parts) >= 2:
            # Vérifier que les parties ne sont pas des nombres
            if not any(part.isdigit() for part in name_parts[:2]):
                result = ' '.join([part.capitalize() for part in name_parts[:2]])
                logger.debug(f"   ✅ Email extraction: {result}")
                return result
    
    return ""

# ================================
# 🧩 FONCTION FIZARANA NOM/PRÉNOM
# ================================
def split_fullname(fullname: str) -> Dict[str, str]:
    """
    Diviser un nom complet en Nom et Prénom(s)
    Gestion des formats malagasy et français
    """
    if not fullname:
        return {"nom": "", "prenom": "", "nom_complet": ""}
    
    words = [w.strip() for w in fullname.strip().split() if w.strip()]
    
    # Si tsy misy anarana na iray monja
    if len(words) == 0:
        return {"nom": "", "prenom": "", "nom_complet": ""}
    elif len(words) == 1:
        return {"nom": words[0], "prenom": "", "nom_complet": fullname}
    
    # Format 1: "RAKOTO Jean" (NOM en majuscule, prénom Capitalize)
    if words[0].isupper() and len(words[0]) > 2:
        nom = words[0]
        prenom = ' '.join(words[1:])
        return {
            "nom": nom,
            "prenom": prenom,
            "nom_complet": fullname
        }
    
    # Format 2: "Jean RAKOTO" (prénom puis NOM)
    if words[-1].isupper() and len(words[-1]) > 2:
        nom = words[-1]
        prenom = ' '.join(words[:-1])
        return {
            "nom": nom,
            "prenom": prenom,
            "nom_complet": fullname
        }
    
    # Format 3: Tous Capitalize - premier mot = prénom, dernier mot = nom (pour noms malagasy)
    if len(words) >= 2:
        # Pour noms malagasy simple: "Jean Rakoto" (2 mots)
        if len(words) == 2:
            # Vérifier si le dernier mot est un nom malagasy courant
            last_word_lower = words[-1].lower()
            if (last_word_lower in MALAGASY_NAMES or 
                any(name in last_word_lower for name in MALAGASY_NAMES)):
                return {"nom": words[-1], "prenom": words[0], "nom_complet": fullname}
            else:
                return {"nom": words[0], "prenom": words[1], "nom_complet": fullname}
        
        # Pour 3 mots: "RANDRIANARIVONY Jean Marie"
        elif len(words) == 3:
            if words[0].isupper():
                return {"nom": words[0], "prenom": ' '.join(words[1:]), "nom_complet": fullname}
            elif words[-1].isupper():
                return {"nom": words[-1], "prenom": ' '.join(words[:-1]), "nom_complet": fullname}
            else:
                # Hypothèse: premier mot = nom, autres = prénoms
                return {"nom": words[0], "prenom": ' '.join(words[1:]), "nom_complet": fullname}
    
    # Fallback: premier mot = nom, reste = prénom
    return {
        "nom": words[0],
        "prenom": ' '.join(words[1:]) if len(words) > 1 else "",
        "nom_complet": fullname
    }

# ================================
# 📞 EXTRACTION TÉLÉPHONE AMÉLIORÉE
# ================================
def extract_phone(text: str) -> str:
    """Extraire téléphone avec formats Madagascar"""
    # Nettoyer le texte
    clean_text = re.sub(r'[^\d\s\+\(\)\-]', '', text)
    
    patterns = [
        # Madagascar: +261 32 123 45 67 ou 032 12 345 67
        r'(?:(?:\+|00)261|0)\s*[2-9]\d?\s*\d{2}\s*\d{2}\s*\d{2}',
        # Madagascar compact: +261321234567
        r'(?:\+|00)261[2-9]\d{8}',
        # France: +33 1 23 45 67 89
        r'(?:(?:\+|00)33|0)[1-9](?:\s*\d{2}){4}',
        # International générique
        r'\+?\d[\d\s\-\(\)]{8,}\d',
        # 10 chiffres consécutifs
        r'\b\d{10}\b',
    ]
    
    phones_found = []
    for pattern in patterns:
        matches = re.findall(pattern, clean_text)
        phones_found.extend(matches)
    
    # Filtrer et prioriser
    valid_phones = []
    for phone in phones_found:
        # Nettoyer
        phone_clean = re.sub(r'\D', '', phone)
        if 9 <= len(phone_clean) <= 13:
            # Formater joliment
            if phone_clean.startswith('261'):
                # Format: +261 XX XXX XX XX
                if len(phone_clean) == 12:  # +261XXXXXXXXX
                    formatted = f"+261 {phone_clean[3:5]} {phone_clean[5:8]} {phone_clean[8:10]} {phone_clean[10:]}"
                else:
                    formatted = f"+261 {phone_clean[3:]}"
            elif phone_clean.startswith('33'):
                formatted = f"+33 {phone_clean[2:]}"
            elif len(phone_clean) == 10:
                formatted = f"{phone_clean[:2]} {phone_clean[2:4]} {phone_clean[4:6]} {phone_clean[6:8]} {phone_clean[8:]}"
            else:
                formatted = phone_clean
            
            if formatted not in valid_phones:
                valid_phones.append(formatted)
    
    return valid_phones[0] if valid_phones else ""

# ================================
# 🎯 EXTRACTION COMPÉTENCES STRUCTURÉES
# ================================
def extract_skills_structured(text: str) -> Dict[str, List[str]]:
    """Extraire compétences par catégorie"""
    text_lower = text.lower()
    skills = {}
    
    for category, keywords in TECH_SKILLS.items():
        found = []
        for keyword in keywords:
            # Recherche exacte ou avec variations
            pattern = r'\b' + re.escape(keyword) + r'(?:s|es)?\b'
            if re.search(pattern, text_lower):
                found.append(keyword.title())
        
        if found:
            skills[category] = list(set(found))
    
    return skills

def extract_experience_years(text: str) -> int:
    """Extraire années d'expérience précise"""
    # Chercher dans différentes formulations
    patterns = [
        r'(\d+)\s*(?:an|ans|année|années|year|years)\s*(?:d\'?expérience|experience|d\'?exp)',
        r'expérience\s*[:=]\s*(\d+)\s*(?:an|ans|année|années)',
        r'(\d+)\s*ans?\s*(?:d\'?exp|expérience)',
        r'(\d+)\+?\s*(?:an|ans)\s*.*expérience',
    ]
    
    max_years = 0
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            try:
                years = int(match)
                max_years = max(max_years, years)
            except:
                pass
    
    # Si pas trouvé, estimer par dates
    if max_years == 0:
        date_pattern = r'(?:19|20)\d{2}'
        dates = re.findall(date_pattern, text)
        if len(dates) >= 2:
            try:
                dates = [int(d) for d in dates if 1900 <= int(d) <= datetime.now().year]
                if dates:
                    oldest = min(dates)
                    current_year = datetime.now().year
                    max_years = max(0, current_year - oldest - 18)  # -18 pour âge études
            except:
                pass
    
    return min(max_years, 40)  # Limiter à 40 ans

# ================================
# 🏫 EXTRACTION FORMATION
# ================================
def extract_education(text: str) -> List[Dict[str, str]]:
    """Extraire formation structurée"""
    education = []
    lines = text.split('\n')
    
    edu_keywords = ['bac', 'licence', 'master', 'doctorat', 'diplôme', 'diploma',
                   'école', 'université', 'university', 'institut', 'bts', 'dut',
                   'ingénieur', 'engineer', 'mba', 'phd', 'doctor', 'formation',
                   'graduat', 'certificat', 'certification', 'brevet', 'niveau']
    
    for i, line in enumerate(lines):
        line_lower = line.lower()
        if any(keyword in line_lower for keyword in edu_keywords):
            # Chercher l'année
            year_match = re.search(r'(?:19|20)\d{2}', line)
            year = year_match.group(0) if year_match else ""
            
            # Chercher l'établissement (mots en majuscules)
            org_match = re.search(r'[A-ZÀ-Ý][A-ZÀ-Ýa-zà-ÿ\s\-&]+', line)
            institution = org_match.group(0).strip() if org_match else ""
            
            education.append({
                "diploma": line.strip(),
                "year": year,
                "institution": institution
            })
    
    return education[:5]  # Limiter à 5 formations

# ================================
# 🚀 FONCTION PRINCIPALE OPTIMISÉE AVEC SPLIT NAME
# ================================
def extract_info(content: bytes, filename: str, sender_name: str = "") -> Dict[str, Any]:
    """Fonction principale - version production améliorée"""
    start_time = datetime.now()
    
    try:
        logger.info(f"🔍 Début extraction: {filename} ({len(content)} bytes)")
        
        # 1. Extraction texte
        text = extract_text_from_bytes(content, filename)
        
        if not text or len(text.strip()) < 50:
            logger.warning("📭 CV trop court ou vide")
            return {
                "confidence": 0,
                "error": "CV vide ou illisible",
                "structured": {},
                "text_length": len(text) if text else 0
            }
        
        logger.info(f"📄 Texte extrait: {len(text)} caractères")
        
        # 2. Détection langue
        language = "fr"
        if LANGDETECT_AVAILABLE:
            try:
                language = detect_func(text[:500])
            except:
                language = "fr"
        
        # 3. Extraction des informations
        fullname = extract_fullname(text)
        
        # 4. Fallback: Utiliser le nom de l'expéditeur si extraction échoue
        if not fullname and sender_name:
            logger.info(f"   📧 Utilisation nom expéditeur: {sender_name}")
            fullname = sender_name
        
        # 5. Split nom/prénom
        name_parts = split_fullname(fullname)
        
        email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        email = email_match.group(0) if email_match else ""
        phone = extract_phone(text)
        skills = extract_skills_structured(text)
        experience_years = extract_experience_years(text)
        education = extract_education(text)
        
        # Calculer le score total de compétences
        total_skills = sum(len(skill_list) for skill_list in skills.values())
        
        # 6. Calcul confiance avec bonus pour nom correct
        confidence = 0
        if fullname: 
            # Vérifier que ce n'est pas un nom de ville
            is_town = any(town in fullname.lower() for town in MALAGASY_TOWN_NAMES)
            if not is_town:
                confidence += 25
                # Bonus si le nom est bien formaté (majuscule/minuscule)
                if re.match(r'^[A-ZÀ-Ý]{2,} [A-ZÀ-Ý][a-zà-ÿ]+', fullname):
                    confidence += 5
        
        if email: confidence += 25
        if phone: confidence += 15
        if total_skills > 0: confidence += min(total_skills * 2, 25)
        if experience_years > 0: confidence += 10
        
        confidence = min(confidence, 100)
        
        # 7. Structuration du résultat AVEC NOM/PRÉNOM SÉPARÉS
        structured = {
            "fullname": fullname,
            "nom": name_parts["nom"],
            "prenom": name_parts["prenom"],
            "email": email,
            "phone": phone,
            "skills": skills,
            "total_skills": total_skills,
            "experience_years": experience_years,
            "education": education,
            "text_preview": text[:300] + "..." if len(text) > 300 else text,
        }
        
        # 8. Logs détaillés
        processing_time = (datetime.now() - start_time).total_seconds()
        logger.info(f"✅ Extraction réussie en {processing_time:.2f}s")
        logger.info(f"   👤 Nom complet: {fullname or 'Non trouvé'}")
        if name_parts["nom"] or name_parts["prenom"]:
            logger.info(f"   📋 Nom: '{name_parts['nom']}', Prénom: '{name_parts['prenom']}'")
        logger.info(f"   📧 {email or 'Non trouvé'}")
        logger.info(f"   📞 {phone or 'Non trouvé'}")
        logger.info(f"   🔧 {total_skills} compétences, {experience_years} ans exp")
        logger.info(f"   💯 Confiance: {confidence}%")
        
        return {
            "language": language,
            "confidence": confidence,
            "structured": structured,
            "raw_text": text[:50000],
            "processing_time": processing_time,
            "parsing_method": "production_v2",
            "timestamp": datetime.now().isoformat(),
            "sender_name_used": bool(sender_name and not fullname)
        }
        
    except Exception as e:
        logger.error(f"❌ Erreur extraction: {e}")
        logger.error(traceback.format_exc())
        
        return {
            "confidence": 0,
            "error": f"Erreur extraction: {str(e)[:100]}",
            "structured": {},
            "raw_text": "",
            "processing_time": (datetime.now() - start_time).total_seconds(),
            "timestamp": datetime.now().isoformat()
        }

# ================================
# 🧪 FONCTION DE TEST AMÉLIORÉE
# ================================
def test_parsing():
    """Fonction de test pour vérifier le parsing"""
    test_cases = [
        ("RAKOTO Jean\nEmail: jean.rakoto@gmail.com\nTél: 032 12 345 67\nFianarantsoa", 
         "RAKOTO Jean", "RAKOTO", "Jean"),
        ("Jean RANDRIANARISOA\n+261 34 56 78 90\nAntananarivo", 
         "Jean RANDRIANARISOA", "RANDRIANARISOA", "Jean"),
        ("Andrianaivo Soanierana\nFianarantsoa Madagascar", 
         "", "", ""),  # Devrait être ignoré (ville)
        ("RASOA Manantsoa\nrue Ankaditapaka", 
         "RASOA Manantsoa", "RASOA", "Manantsoa"),
        ("Njaka RAKOTONDRABE\nnjaka.rakotondrabe@email.com", 
         "Njaka RAKOTONDRABE", "RAKOTONDRABE", "Njaka"),
        ("Marie-Claude DUPONT\nParis France", 
         "Marie-Claude DUPONT", "DUPONT", "Marie-Claude"),
    ]
    
    print("🧪 TEST PARSING AVEC FILTRE TANÀNA")
    print("=" * 60)
    
    for i, (text, expected_fullname, expected_nom, expected_prenom) in enumerate(test_cases):
        result = extract_info(text.encode('utf-8'), f"test_{i}.txt")
        structured = result.get('structured', {})
        
        print(f"\nTest {i+1}: {text[:50]}...")
        print(f"  Nom complet: '{structured.get('fullname', '')}' (attendu: '{expected_fullname}')")
        print(f"  Nom: '{structured.get('nom', '')}' (attendu: '{expected_nom}')")
        print(f"  Prénom: '{structured.get('prenom', '')}' (attendu: '{expected_prenom}')")
        print(f"  Confiance: {result.get('confidence', 0)}%")
        
        # Vérification
        fullname_match = structured.get('fullname', '') == expected_fullname
        nom_match = structured.get('nom', '') == expected_nom
        prenom_match = structured.get('prenom', '') == expected_prenom
        
        if fullname_match and nom_match and prenom_match:
            print("  ✅ CORRECT")
        else:
            print("  ❌ INCORRECT")

# ================================
# 📁 FONCTION ALTERNATIVE POUR TEXTE DÉJÀ EXTRAIT
# ================================
def extract_info_from_text(text: str, filename: str = "cv_unknown.txt", sender_name: str = "") -> Dict[str, Any]:
    """Version alternative pour extraction depuis texte déjà extrait"""
    try:
        # Convertir texte en bytes pour compatibilité
        content = text.encode('utf-8')
        return extract_info(content, filename, sender_name)
    except Exception as e:
        logger.error(f"Erreur extract_info_from_text: {e}")
        return {
            "confidence": 30, 
            "error": str(e),
            "structured": {
                "fullname": "",
                "nom": "",
                "prenom": "",
                "email": "",
                "phone": "",
                "skills": {},
                "total_skills": 0,
                "experience_years": 0,
                "education": []
            }
        }

# ================================
# 🏃 EXÉCUTION TEST SI FICHIER PRINCIPAL
# ================================
if __name__ == "__main__":
    print("\n" + "="*60)
    print("🧠 TEST PARSING CV AUTOMATIQUE")
    print("="*60)
    
    # Configuration logging pour test
    logging.basicConfig(level=logging.INFO)
    
    # Exécuter les tests
    test_parsing()
    
    print("\n" + "="*60)
    print("✅ TEST TERMINÉ")
    print("="*60)
    
    
    
    
    
    # Ao amin'ny faran'ny parsing.py - manampy fonction backup
def extract_info_from_text(text: str, filename: str = "cv_unknown.txt") -> Dict[str, Any]:
    """Version alternative pour extraction depuis texte deja extrait"""
    try:
        # Convertir texte en bytes pour compatibilité
        content = text.encode('utf-8')
        return extract_info(content, filename)
    except Exception as e:
        logger.error(f"Erreur extract_info_from_text: {e}")
        return {"confidence": 30, "error": str(e)} 