

# import os
# from io import BytesIO
# from typing import List, Optional

# from fastapi import UploadFile
# from sqlalchemy.orm import Session
# from sqlalchemy import text

# from minio import Minio
# from minio.error import S3Error

# from app.services.parsing import extract_info
# from app.services.scoring_auto import calculer_score_auto

# # ==========================================================
# # 🔹 Configuration MinIO
# # ==========================================================
# MINIO_ENDPOINT = os.getenv("MINIO_ENDPOINT", "127.0.0.1:9000")
# MINIO_ACCESS_KEY = os.getenv("MINIO_ROOT_USER", "jeremi")
# MINIO_SECRET_KEY = os.getenv("MINIO_ROOT_PASSWORD", "Jeremi123")
# MINIO_BUCKET = os.getenv("MINIO_BUCKET", "cvs")

# minio_client = Minio(
#     MINIO_ENDPOINT,
#     access_key=MINIO_ACCESS_KEY,
#     secret_key=MINIO_SECRET_KEY,
#     secure=False
# )

# if not minio_client.bucket_exists(MINIO_BUCKET):
#     minio_client.make_bucket(MINIO_BUCKET)
#     print(f"🔹 Bucket MinIO créé: {MINIO_BUCKET}")
# else:
#     print(f"🔹 Bucket MinIO existe déjà: {MINIO_BUCKET}")

# # ==========================================================
# # 🔹 Fonction principale
# # ==========================================================
# def save_upload_file(
#     db: Session,
#     file: UploadFile,
#     candidature_id: int,
#     offre: Optional[dict] = None,
#     is_cv: bool = True
# ) -> dict:
#     """
#     - Upload fichier vers MinIO
#     - Enregistre le chemin dans raw_cv_s3 (UNIQUEMENT pour le CV)
#     - Parse le CV
#     - Extrait les infos
#     - Calcule le score automatiquement
#     """

#     if not file:
#         raise ValueError("Aucun fichier fourni")

#     filename = file.filename
#     content = file.file.read()
#     file.file.seek(0)

#     print(f"📤 Upload fichier: {filename} ({len(content)} bytes)")

#     # ======================================================
#     # 🔹 Chemin MinIO UNIQUE par candidature
#     # ======================================================
#     object_name = f"cv/{candidature_id}/{filename}"

#     # ======================================================
#     # 🔹 Upload MinIO (BytesIO obligatoire)
#     # ======================================================
#     try:
#         minio_client.put_object(
#             bucket_name=MINIO_BUCKET,
#             object_name=object_name,
#             data=BytesIO(content),
#             length=len(content),
#             content_type=file.content_type
#         )
#         print(f"✅ MinIO OK: {MINIO_BUCKET}/{object_name}")
#     except S3Error as e:
#         print(f"❌ Erreur MinIO: {e}")
#         raise Exception("Erreur upload MinIO")

#     # ======================================================
#     # 🔹 Update DB raw_cv_s3 (CV SEULEMENT)
#     # ======================================================
#     if is_cv:
#         try:
#             db.execute(
#                 text("""
#                     UPDATE candidatures
#                     SET raw_cv_s3 = :path
#                     WHERE id = :id
#                 """),
#                 {
#                     "path": object_name,
#                     "id": candidature_id
#                 }
#             )
#             db.commit()
#             print(f"✅ DB raw_cv_s3 mis à jour (ID={candidature_id})")
#         except Exception as e:
#             db.rollback()
#             print(f"❌ Erreur DB raw_cv_s3: {e}")

#     # ======================================================
#     # 🔹 Parsing du fichier
#     # ======================================================
#     text_content = ""

#     try:
#         if filename.lower().endswith(".pdf"):
#             from app.services.parsing import parse_pdf
#             text_content = parse_pdf(BytesIO(content))
#         elif filename.lower().endswith(".docx"):
#             from app.services.parsing import parse_docx
#             text_content = parse_docx(BytesIO(content))
#     except Exception as e:
#         print(f"❌ Erreur parsing fichier: {e}")

#     print(f"✏️ Texte extrait: {len(text_content)} caractères")

#     # ======================================================
#     # 🔹 Extraction d'informations
#     # ======================================================
#     info = {}
#     if text_content:
#         try:
#             info = extract_info(text_content)
#         except Exception as e:
#             print(f"❌ Erreur extract_info: {e}")

#     # ======================================================
#     # 🔹 Calcul automatique du score
#     # ======================================================
#     score_result = {"score": 0, "passed_threshold": False}

#     if text_content and offre:
#         try:
#             score_result = calculer_score_auto(text_content, offre)
#         except Exception as e:
#             print(f"❌ Erreur scoring: {e}")

#     return {
#         "filename": filename,
#         "object_name": object_name,
#         "text_preview": text_content[:1000],
#         "firstname": info.get("firstname"),
#         "lastname": info.get("lastname"),
#         "score": score_result["score"],
#         "passed_threshold": score_result["passed_threshold"]
#     }

# # ==========================================================
# # 🔹 Traitement depuis bytes (mail)
# # ==========================================================
# def process_cv_from_bytes(
#     db: Session,
#     file_bytes: bytes,
#     filename: str,
#     candidature_id: int,
#     offre: Optional[dict] = None
# ) -> dict:
#     fake_file = UploadFile(
#         filename=filename,
#         file=BytesIO(file_bytes)
#     )
#     return save_upload_file(
#         db=db,
#         file=fake_file,
#         candidature_id=candidature_id,
#         offre=offre,
#         is_cv=True
#     )

# # ==========================================================
# # 🔹 Formulaire (CV + autres fichiers)
# # ==========================================================
# def save_formulaire_files(
#     db: Session,
#     cv: Optional[UploadFile],
#     lm: Optional[UploadFile],
#     diplomes: Optional[UploadFile],
#     candidature_id: int,
#     offre: Optional[dict] = None
# ) -> dict:
#     results = {}

#     if cv:
#         results["cv"] = save_upload_file(
#             db=db,
#             file=cv,
#             candidature_id=candidature_id,
#             offre=offre,
#             is_cv=True
#         )

#     if lm:
#         results["lm"] = save_upload_file(
#             db=db,
#             file=lm,
#             candidature_id=candidature_id,
#             offre=None,
#             is_cv=False
#         )

#     if diplomes:
#         results["diplomes"] = save_upload_file(
#             db=db,
#             file=diplomes,
#             candidature_id=candidature_id,
#             offre=None,
#             is_cv=False
#         )

#     return results

# # Alias rétro-compatibilité
# upload_files_to_minio = save_formulaire_files

# # ==========================================================
# # 🔹 Suppression fichier MinIO
# # ==========================================================
# def delete_file(object_name: str):
#     try:
#         if minio_client.bucket_exists(MINIO_BUCKET):
#             minio_client.remove_object(MINIO_BUCKET, object_name)
#             print(f"🗑️ Fichier supprimé: {object_name}")
#     except S3Error as e:
#         print(f"❌ Erreur suppression MinIO: {e}")
#         raise Exception("Erreur suppression MinIO")









# # app/services/upload_service.py - VERSION CORRIGÉE COMPLÈTE
# import os, io, logging, re
# from datetime import datetime
# from typing import Optional, Dict, Any, List, Tuple
# import uuid
# from pathlib import Path
# import tempfile

# logger = logging.getLogger(__name__)

# # ==================== SERVICE MINIO ====================

# class MinioService:
#     """Service pour l'upload vers MinIO"""
    
#     def __init__(self):
#         try:
#             from minio import Minio
            
#             # Configuration
#             self.endpoint = os.getenv("MINIO_ENDPOINT", "localhost:9000")
#             self.access_key = os.getenv("MINIO_ACCESS_KEY", "jeremi")
#             self.secret_key = os.getenv("MINIO_SECRET_KEY", "Jeremi123")
#             self.bucket_name = os.getenv("MINIO_BUCKET", "siirh-candidatures")
            
#             logger.info(f"🔗 Initialisation MinIO: {self.endpoint}")
            
#             # Client MinIO
#             self.minio_client = Minio(
#                 endpoint=self.endpoint,
#                 access_key=self.access_key,
#                 secret_key=self.secret_key,
#                 secure=False
#             )
            
#             # Vérifier/créer bucket
#             self._ensure_bucket()
#             self.minio_available = True
#             logger.info("✅ Service MinIO initialisé avec succès")
            
#         except Exception as e:
#             logger.error(f"❌ Erreur initialisation MinIO: {e}")
#             self.minio_available = False
#             self.minio_client = None
    
#     def _ensure_bucket(self):
#         """Créer le bucket s'il n'existe pas"""
#         try:
#             if not self.minio_client.bucket_exists(self.bucket_name):
#                 self.minio_client.make_bucket(self.bucket_name)
#                 logger.info(f"✅ Bucket créé: {self.bucket_name}")
#             else:
#                 logger.info(f"✅ Bucket existe déjà: {self.bucket_name}")
#         except Exception as e:
#             logger.error(f"❌ Erreur création bucket: {e}")
#             raise
    
#     def upload_cv(self, file_data: bytes, filename: str, offre_ref: str, candidate_email: str) -> Optional[str]:
#         """
#         Uploader un CV vers MinIO
#         Structure: offres/{ref}/candidats/{email}/cv_{timestamp}_{id}.pdf
#         """
#         if not self.minio_available or not self.minio_client:
#             logger.error("❌ Service MinIO non disponible")
#             return None
        
#         try:
#             # Nettoyer les noms pour le chemin
#             safe_ref = re.sub(r'[^\w\-]', '_', offre_ref)
#             email_local = candidate_email.split('@')[0]
#             safe_email = re.sub(r'[^\w\-]', '_', email_local)[:50]
            
#             # Extension fichier
#             file_ext = Path(filename).suffix.lower()
#             if not file_ext:
#                 file_ext = '.pdf'
            
#             # Nom unique
#             timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#             unique_id = str(uuid.uuid4())[:8]
#             new_filename = f"cv_{timestamp}_{unique_id}{file_ext}"
            
#             # Chemin MinIO
#             object_path = f"offres/{safe_ref}/candidats/{safe_email}/{new_filename}"
            
#             # Content-type
#             content_type = self._get_content_type(file_ext)
            
#             # Upload
#             self.minio_client.put_object(
#                 bucket_name=self.bucket_name,
#                 object_name=object_path,
#                 data=io.BytesIO(file_data),
#                 length=len(file_data),
#                 content_type=content_type
#             )
            
#             logger.info(f"✅ CV uploadé vers MinIO: {object_path} ({len(file_data)} bytes)")
#             return object_path
            
#         except Exception as e:
#             logger.error(f"❌ Erreur upload MinIO: {e}")
#             return None
    
#     def _get_content_type(self, extension: str) -> str:
#         """Déterminer le type MIME"""
#         types = {
#             '.pdf': 'application/pdf',
#             '.doc': 'application/msword',
#             '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
#             '.odt': 'application/vnd.oasis.opendocument.text',
#             '.txt': 'text/plain',
#             '.rtf': 'application/rtf',
#         }
#         return types.get(extension.lower(), 'application/octet-stream')

# # Instance globale
# minio_service = MinioService()

# # ==================== FONCTIONS D'EXTRACTION TEXTE ====================

# def extract_text_from_bytes(content: bytes, filename: str) -> str:
#     """
#     Extraire le texte d'un fichier CV selon son format
#     Support: PDF, DOCX, DOC, TXT
#     """
#     filename_lower = filename.lower()
    
#     # PDF
#     if filename_lower.endswith('.pdf'):
#         try:
#             import PyPDF2
#             pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
#             text = ""
            
#             # Extraire les 3 premières pages max
#             for i in range(min(3, len(pdf_reader.pages))):
#                 page = pdf_reader.pages[i]
#                 page_text = page.extract_text()
#                 if page_text:
#                     text += page_text + "\n"
            
#             return text.strip() if text else ""
#         except Exception as e:
#             logger.warning(f"⚠️ Erreur extraction PDF: {e}")
    
#     # DOCX
#     elif filename_lower.endswith('.docx'):
#         try:
#             from docx import Document
#             doc = Document(io.BytesIO(content))
#             paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
#             return "\n".join(paragraphs[:100])  # 100 paragraphes max
#         except Exception as e:
#             logger.warning(f"⚠️ Erreur extraction DOCX: {e}")
    
#     # DOC (ancien format)
#     elif filename_lower.endswith('.doc'):
#         try:
#             # Essayer comme texte brut
#             text = content.decode('utf-8', errors='ignore')
#             if len(text) > 100:
#                 return text[:3000]
#         except:
#             try:
#                 text = content.decode('latin-1', errors='ignore')
#                 if len(text) > 100:
#                     return text[:3000]
#             except:
#                 pass
    
#     # Texte brut
#     try:
#         text = content.decode('utf-8', errors='ignore')
#         if len(text) > 100:
#             return text[:3000]
#     except:
#         try:
#             text = content.decode('latin-1', errors='ignore')
#             if len(text) > 100:
#                 return text[:3000]
#         except:
#             pass
    
#     return ""

# # ==================== FONCTION PRINCIPALE ====================

# def process_cv_from_bytes(db, content: bytes, filename: str, candidature_id: int) -> Dict[str, Any]:
#     """
#     Traitement complet d'un CV:
#     1. Upload vers MinIO
#     2. Extraction texte
#     3. Parsing NLP
#     4. Calcul score
#     5. Mise à jour candidature
#     """
#     try:
#         from app.models.models import Candidature
#         from app.models.offres import Offre
#         from app.services.parsing import extract_info
#         from app.services.scoring_auto import calculate_cv_score
        
#         logger.info(f"📤 Début traitement CV: {filename} ({len(content)} bytes)")
        
#         # Récupérer candidature
#         candidature = db.query(Candidature).filter(Candidature.id == candidature_id).first()
#         if not candidature:
#             logger.error(f"❌ Candidature {candidature_id} non trouvée")
#             return {"success": False, "error": "Candidature non trouvée"}
        
#         # Récupérer offre associée
#         offre = None
#         if candidature.offre_id:
#             offre = db.query(Offre).filter(Offre.id == candidature.offre_id).first()
        
#         # Référence offre pour chemin MinIO
#         offre_ref = getattr(candidature, 'ref_offre', 'UNASSIGNED')
        
#         logger.info(f"   📋 Offre référence: {offre_ref}")
#         logger.info(f"   👤 Candidat: {candidature.email}")
        
#         # 🔥 ÉTAPE 1: UPLOAD VERS MINIO
#         minio_path = minio_service.upload_cv(
#             file_data=content,
#             filename=filename,
#             offre_ref=offre_ref,
#             candidate_email=candidature.email
#         )
        
#         if not minio_path:
#             logger.error(f"❌ Échec upload MinIO")
#             return {"success": False, "error": "Échec upload MinIO"}
        
#         # Sauvegarder chemin MinIO
#         candidature.raw_cv_s3 = minio_path
        
#         # 🔥 ÉTAPE 2: EXTRACTION TEXTE
#         cv_text = ""
#         nlp_info = {}
#         score_final = 50  # Score par défaut
        
#         try:
#             cv_text = extract_text_from_bytes(content, filename)
            
#             if cv_text and len(cv_text) > 100:
#                 logger.info(f"   📝 Texte extrait: {len(cv_text)} caractères")
                
#                 # 🔥 ÉTAPE 3: EXTRACTION NLP - CORRECTION ICI
#                 try:
#                     # 🔴 CORRECTION: Passer content ET filename comme arguments
#                     nlp_info = extract_info(content, filename)  # ← CORRIGÉ
#                     logger.info(f"   ✅ Extraction NLP terminée")
                    
#                     # Score d'extraction (utiliser 'confidence' depuis nlp_info)
#                     score_extraction = nlp_info.get('confidence', 30)
#                     logger.info(f"   📊 Score extraction NLP: {score_extraction}%")
                    
#                 except Exception as e:
#                     logger.error(f"   ❌ Erreur extraction NLP: {e}")
#                     nlp_info = {
#                         'confidence': 30,
#                         'structured': {
#                             'fullname': None,
#                             'phone': None,
#                             'email': None,
#                             'skills': {},
#                             'experience_years': 0,
#                             'education': []
#                         },
#                         'error': str(e)
#                     }
#                     score_extraction = 30
                
#                 # 🔥 ÉTAPE 4: CALCUL SCORE MATCHING AVEC OFFRE
#                 score_matching = 50
#                 if offre and hasattr(offre, 'mission') and offre.mission:
#                     try:
#                         score_result = calculate_cv_score(cv_text, offre.mission)
#                         if score_result and 'score_total' in score_result:
#                             score_matching = score_result['score_total']
#                             logger.info(f"   🎯 Score matching offre: {score_matching}%")
#                     except Exception as e:
#                         logger.error(f"   ❌ Erreur calcul score matching: {e}")
                
#                 # 🔥 ÉTAPE 5: SCORE FINAL (moyenne pondérée)
#                 score_final = int(score_extraction * 0.4 + score_matching * 0.6)
#                 logger.info(f"   🎯 Score final: {score_final}% (Extraction: {score_extraction}%, Matching: {score_matching}%)")
                
#                 # 🔥 ÉTAPE 6: PRÉPARATION DONNÉES PERSISTANCE
#                 parsed_data = {
#                     'extraction_date': datetime.now().isoformat(),
#                     'source_file': filename,
#                     'text_length': len(cv_text),
#                     'score_extraction': score_extraction,
#                     'score_matching': score_matching,
#                     'score_final': score_final,
#                     'extracted_info': nlp_info
#                 }
                
#                 # Mettre à jour candidature
#                 candidature.parsed_json = parsed_data
#                 candidature.score = float(score_final)
#                 candidature.cv_text = cv_text[:2000]  # Stocker échantillon
                
#                 # Mettre à jour informations extraites depuis NLP
#                 if nlp_info and 'structured' in nlp_info:
#                     structured = nlp_info.get('structured', {})
                    
#                     # Mettre à jour nom complet
#                     if structured.get('fullname'):
#                         candidature.fullname = structured['fullname']
#                         logger.info(f"   👤 Nom extrait: {structured['fullname']}")
                    
#                     # Mettre à jour téléphone
#                     if structured.get('phone'):
#                         candidature.phone = structured['phone']
#                         logger.info(f"   📞 Téléphone extrait: {structured['phone']}")
                    
#                     # Mettre à jour email
#                     if structured.get('email') and not candidature.email:
#                         candidature.email = structured['email']
                    
#                     # Mettre à jour compétences
#                     if structured.get('skills'):
#                         # Convertir le dictionnaire de compétences en liste
#                         all_skills = []
#                         for category, skills in structured['skills'].items():
#                             if isinstance(skills, list):
#                                 all_skills.extend(skills)
                        
#                         if all_skills:
#                             candidature.competences = ', '.join(all_skills[:10])  # Limiter à 10 compétences
#                             logger.info(f"   🔧 Compétences extraites: {len(all_skills)}")
                    
#                     # Mettre à jour années d'expérience
#                     if structured.get('experience_years'):
#                         candidature.experience_years = structured['experience_years']
#                         logger.info(f"   📅 Expérience: {structured['experience_years']} ans")
                
#                 # Stocker données NLP complètes
#                 candidature.nlp_data = nlp_info
                
#             else:
#                 logger.warning(f"   ⚠️ Texte CV trop court ou impossible à extraire")
#                 score_final = 10
#                 candidature.score = 10.0
#                 nlp_info = {'confidence': 10, 'error': 'Texte trop court'}
                
#         except Exception as e:
#             logger.error(f"   ❌ Erreur traitement texte CV: {e}")
#             score_final = 10
#             candidature.score = 10.0
#             nlp_info = {'confidence': 10, 'error': str(e)}
        
#         # 🔥 ÉTAPE 7: SAUVEGARDE BASE DE DONNÉES
#         db.commit()
#         logger.info(f"✅ Candidature mise à jour avec succès")
        
#         # 🔥 RETOUR RÉSULTATS
#         return {
#             "success": True,
#             "minio_path": minio_path,
#             "candidature_id": candidature_id,
#             "filename": filename,
#             "score": score_final,
#             "nlp_info": nlp_info,
#             "message": "CV traité avec succès"
#         }
            
#     except Exception as e:
#         logger.error(f"❌ Erreur traitement CV: {e}")
#         import traceback
#         logger.error(traceback.format_exc())
#         return {"success": False, "error": str(e)}

# def upload_files_to_minio(files: List[Tuple[str, bytes]], offre_ref: str, candidate_info: Dict[str, Any]) -> Dict[str, Any]:
#     """Upload multiple de fichiers vers MinIO (pour batch processing)"""
#     results = {
#         "success": [],
#         "failed": [],
#         "total_files": len(files)
#     }
    
#     for filename, file_data in files:
#         try:
#             minio_path = minio_service.upload_cv(
#                 file_data=file_data,
#                 filename=filename,
#                 offre_ref=offre_ref,
#                 candidate_email=candidate_info.get('email', 'unknown@example.com')
#             )
            
#             if minio_path:
#                 results["success"].append({
#                     "filename": filename,
#                     "minio_path": minio_path,
#                     "size_bytes": len(file_data)
#                 })
#                 logger.info(f"✅ Upload réussi: {filename}")
#             else:
#                 results["failed"].append({
#                     "filename": filename,
#                     "error": "Échec upload MinIO"
#                 })
#                 logger.error(f"❌ Échec upload: {filename}")
                
#         except Exception as e:
#             results["failed"].append({
#                 "filename": filename,
#                 "error": str(e)
#             })
#             logger.error(f"❌ Erreur upload {filename}: {e}")
    
#     results["success_count"] = len(results["success"])
#     results["failed_count"] = len(results["failed"])
    
#     return results

# # ==================== FONCTIONS UTILITAIRES ====================

# def test_minio_connection() -> Dict[str, Any]:
#     """Tester la connexion MinIO"""
#     if not minio_service.minio_available:
#         return {
#             "status": "error",
#             "message": "Service MinIO non disponible",
#             "endpoint": minio_service.endpoint
#         }
    
#     try:
#         # Créer un fichier test
#         test_content = b"Test de connexion MinIO - CV de test"
#         test_filename = "test_cv.pdf"
        
#         minio_path = minio_service.upload_cv(
#             file_data=test_content,
#             filename=test_filename,
#             offre_ref="TEST_CONNECTION",
#             candidate_email="test@example.com"
#         )
        
#         if minio_path:
#             return {
#                 "status": "success",
#                 "message": "Connexion MinIO fonctionnelle",
#                 "test_file": minio_path,
#                 "endpoint": minio_service.endpoint,
#                 "bucket": minio_service.bucket_name
#             }
#         else:
#             return {
#                 "status": "warning",
#                 "message": "MinIO accessible mais upload échoué",
#                 "endpoint": minio_service.endpoint
#             }
            
#     except Exception as e:
#         return {
#             "status": "error",
#             "message": f"Erreur test MinIO: {str(e)}",
#             "endpoint": minio_service.endpoint
#         }

# def get_file_info(file_data: bytes, filename: str) -> Dict[str, Any]:
#     """Obtenir des informations sur un fichier"""
#     return {
#         "filename": filename,
#         "size_bytes": len(file_data),
#         "extension": Path(filename).suffix.lower(),
#         "is_cv": any(filename.lower().endswith(ext) for ext in ['.pdf', '.doc', '.docx', '.odt']),
#         "estimated_pages": len(file_data) // 5000 if filename.lower().endswith('.pdf') else None
#     }

# # Ajouter à la fin de app/services/upload_service.py

# def save_upload_file(upload_file, destination: Path) -> str:
#     """
#     Fonction de compatibilité pour sauvegarder un fichier uploadé localement
#     Utilisée par d'autres parties du système
#     """
#     import shutil
#     import logging
    
#     logger = logging.getLogger(__name__)
    
#     try:
#         # Créer le répertoire parent si nécessaire
#         destination.parent.mkdir(parents=True, exist_ok=True)
        
#         # Sauvegarder le fichier
#         with open(destination, "wb") as buffer:
#             shutil.copyfileobj(upload_file.file, buffer)
        
#         logger.info(f"📁 Fichier sauvegardé localement: {destination}")
#         return str(destination)
        
#     except Exception as e:
#         logger.error(f"❌ Erreur sauvegarde fichier {destination}: {e}")
#         raise







# app/services/upload_service.py - VERSION CORRIGÉE COMPLÈTE
import os, io, logging, re
from datetime import datetime
from typing import Optional, Dict, Any, List, Tuple
import uuid
from pathlib import Path
import tempfile

logger = logging.getLogger(__name__)

# ==================== SERVICE MINIO ====================

class MinioService:
    """Service pour l'upload vers MinIO"""
    
    def __init__(self):
        try:
            from minio import Minio
            
            # Configuration
            self.endpoint = os.getenv("MINIO_ENDPOINT", "localhost:9000")
            self.access_key = os.getenv("MINIO_ACCESS_KEY", "jeremi")
            self.secret_key = os.getenv("MINIO_SECRET_KEY", "Jeremi123")
            self.bucket_name = os.getenv("MINIO_BUCKET", "siirh-candidatures")
            
            logger.info(f"🔗 Initialisation MinIO: {self.endpoint}")
            
            # Client MinIO
            self.minio_client = Minio(
                endpoint=self.endpoint,
                access_key=self.access_key,
                secret_key=self.secret_key,
                secure=False
            )
            
            # Vérifier/créer bucket
            self._ensure_bucket()
            self.minio_available = True
            logger.info("✅ Service MinIO initialisé avec succès")
            
        except Exception as e:
            logger.error(f"❌ Erreur initialisation MinIO: {e}")
            self.minio_available = False
            self.minio_client = None
    
    def _ensure_bucket(self):
        """Créer le bucket s'il n'existe pas"""
        try:
            if not self.minio_client.bucket_exists(self.bucket_name):
                self.minio_client.make_bucket(self.bucket_name)
                logger.info(f"✅ Bucket créé: {self.bucket_name}")
            else:
                logger.info(f"✅ Bucket existe déjà: {self.bucket_name}")
        except Exception as e:
            logger.error(f"❌ Erreur création bucket: {e}")
            raise
    
    def upload_cv(self, file_data: bytes, filename: str, offre_ref: str, candidate_email: str) -> Optional[str]:
        """
        Uploader un CV vers MinIO
        Structure: offres/{ref}/candidats/{email}/cv_{timestamp}_{id}.pdf
        """
        if not self.minio_available or not self.minio_client:
            logger.error("❌ Service MinIO non disponible")
            return None
        
        try:
            # Nettoyer les noms pour le chemin
            safe_ref = re.sub(r'[^\w\-]', '_', offre_ref)
            email_local = candidate_email.split('@')[0]
            safe_email = re.sub(r'[^\w\-]', '_', email_local)[:50]
            
            # Extension fichier
            file_ext = Path(filename).suffix.lower()
            if not file_ext:
                file_ext = '.pdf'
            
            # Nom unique
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_id = str(uuid.uuid4())[:8]
            new_filename = f"cv_{timestamp}_{unique_id}{file_ext}"
            
            # Chemin MinIO
            object_path = f"offres/{safe_ref}/candidats/{safe_email}/{new_filename}"
            
            # Content-type
            content_type = self._get_content_type(file_ext)
            
            # Upload
            self.minio_client.put_object(
                bucket_name=self.bucket_name,
                object_name=object_path,
                data=io.BytesIO(file_data),
                length=len(file_data),
                content_type=content_type
            )
            
            logger.info(f"✅ CV uploadé vers MinIO: {object_path} ({len(file_data)} bytes)")
            return object_path
            
        except Exception as e:
            logger.error(f"❌ Erreur upload MinIO: {e}")
            return None
    
    def _get_content_type(self, extension: str) -> str:
        """Déterminer le type MIME"""
        types = {
            '.pdf': 'application/pdf',
            '.doc': 'application/msword',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.odt': 'application/vnd.oasis.opendocument.text',
            '.txt': 'text/plain',
            '.rtf': 'application/rtf',
        }
        return types.get(extension.lower(), 'application/octet-stream')

# Instance globale
minio_service = MinioService()

# ==================== FONCTIONS D'EXTRACTION TEXTE ====================

def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """
    Extraire le texte d'un fichier CV selon son format
    Support: PDF, DOCX, DOC, TXT
    """
    filename_lower = filename.lower()
    
    # PDF
    if filename_lower.endswith('.pdf'):
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            
            # Extraire les 3 premières pages max
            for i in range(min(3, len(pdf_reader.pages))):
                page = pdf_reader.pages[i]
                page_text = page.extract_text()
                if page_text:
                    # 🔴 CORRECTION IMPORTANTE: Nettoyer les caractères NULL
                    text += page_text.replace('\x00', ' ').replace('\0', ' ') + "\n"
            
            return text.strip() if text else ""
        except Exception as e:
            logger.warning(f"⚠️ Erreur extraction PDF: {e}")
    
    # DOCX
    elif filename_lower.endswith('.docx'):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n".join(paragraphs[:100])  # 100 paragraphes max
            
            # 🔴 CORRECTION: Nettoyer les caractères NULL
            text = text.replace('\x00', ' ').replace('\0', ' ')
            return text
        except Exception as e:
            logger.warning(f"⚠️ Erreur extraction DOCX: {e}")
    
    # DOC (ancien format)
    elif filename_lower.endswith('.doc'):
        try:
            # Essayer comme texte brut
            text = content.decode('utf-8', errors='ignore')
            if len(text) > 100:
                # 🔴 CORRECTION: Nettoyer les caractères NULL
                text = text.replace('\x00', ' ').replace('\0', ' ')
                return text[:3000]
        except:
            try:
                text = content.decode('latin-1', errors='ignore')
                if len(text) > 100:
                    # 🔴 CORRECTION: Nettoyer les caractères NULL
                    text = text.replace('\x00', ' ').replace('\0', ' ')
                    return text[:3000]
            except:
                pass
    
    # Texte brut
    try:
        text = content.decode('utf-8', errors='ignore')
        if len(text) > 100:
            # 🔴 CORRECTION: Nettoyer les caractères NULL
            text = text.replace('\x00', ' ').replace('\0', ' ')
            return text[:3000]
    except:
        try:
            text = content.decode('latin-1', errors='ignore')
            if len(text) > 100:
                # 🔴 CORRECTION: Nettoyer les caractères NULL
                text = text.replace('\x00', ' ').replace('\0', ' ')
                return text[:3000]
        except:
            pass
    
    return ""

# ==================== FONCTION PRINCIPALE ====================

def process_cv_from_bytes(db, content: bytes, filename: str, candidature_id: int) -> Dict[str, Any]:
    """
    Traitement complet d'un CV:
    1. Upload vers MinIO
    2. Extraction texte
    3. Parsing NLP
    4. Calcul score
    5. Mise à jour candidature
    """
    try:
        from app.models.models import Candidature
        from app.models.offres import Offre
        from app.services.parsing import extract_info
        from app.services.scoring_auto import calculate_cv_score
        
        logger.info(f"📤 Début traitement CV: {filename} ({len(content)} bytes)")
        
        # Récupérer candidature
        candidature = db.query(Candidature).filter(Candidature.id == candidature_id).first()
        if not candidature:
            logger.error(f"❌ Candidature {candidature_id} non trouvée")
            return {"success": False, "error": "Candidature non trouvée"}
        
        # Récupérer offre associée
        offre = None
        if candidature.offre_id:
            offre = db.query(Offre).filter(Offre.id == candidature.offre_id).first()
        
        # Référence offre pour chemin MinIO
        offre_ref = getattr(candidature, 'ref_offre', 'UNASSIGNED')
        
        logger.info(f"   📋 Offre référence: {offre_ref}")
        logger.info(f"   👤 Candidat: {candidature.email}")
        
        # 🔥 ÉTAPE 1: UPLOAD VERS MINIO
        minio_path = minio_service.upload_cv(
            file_data=content,
            filename=filename,
            offre_ref=offre_ref,
            candidate_email=candidature.email
        )
        
        if not minio_path:
            logger.error(f"❌ Échec upload MinIO")
            return {"success": False, "error": "Échec upload MinIO"}
        
        # Sauvegarder chemin MinIO
        candidature.raw_cv_s3 = minio_path
        
        # 🔥 ÉTAPE 2: EXTRACTION TEXTE
        cv_text = ""
        nlp_info = {}
        score_final = 50  # Score par défaut
        
        try:
            cv_text = extract_text_from_bytes(content, filename)
            
            if cv_text and len(cv_text) > 100:
                logger.info(f"   📝 Texte extrait: {len(cv_text)} caractères")
                
                # 🔴 CORRECTION SUPPLEMENTAIRE: Nettoyer tous les caractères NULL du texte complet
                cv_text = cv_text.replace('\x00', ' ').replace('\0', ' ')
                
                # 🔥 ÉTAPE 3: EXTRACTION NLP
                try:
                    nlp_info = extract_info(content, filename)
                    logger.info(f"   ✅ Extraction NLP terminée")
                    
                    # Score d'extraction (utiliser 'confidence' depuis nlp_info)
                    score_extraction = nlp_info.get('confidence', 30)
                    logger.info(f"   📊 Score extraction NLP: {score_extraction}%")
                    
                except Exception as e:
                    logger.error(f"   ❌ Erreur extraction NLP: {e}")
                    nlp_info = {
                        'confidence': 30,
                        'structured': {
                            'fullname': None,
                            'phone': None,
                            'email': None,
                            'skills': {},
                            'experience_years': 0,
                            'education': []
                        },
                        'error': str(e)
                    }
                    score_extraction = 30
                
                # 🔥 ÉTAPE 4: CALCUL SCORE MATCHING AVEC OFFRE
                score_matching = 50
                if offre and hasattr(offre, 'mission') and offre.mission:
                    try:
                        score_result = calculate_cv_score(cv_text, offre.mission)
                        if score_result and 'score_total' in score_result:
                            score_matching = score_result['score_total']
                            logger.info(f"   🎯 Score matching offre: {score_matching}%")
                    except Exception as e:
                        logger.error(f"   ❌ Erreur calcul score matching: {e}")
                
                # 🔥 ÉTAPE 5: SCORE FINAL (moyenne pondérée)
                score_final = int(score_extraction * 0.4 + score_matching * 0.6)
                logger.info(f"   🎯 Score final: {score_final}% (Extraction: {score_extraction}%, Matching: {score_matching}%)")
                
                # 🔥 ÉTAPE 6: PRÉPARATION DONNÉES PERSISTANCE
                parsed_data = {
                    'extraction_date': datetime.now().isoformat(),
                    'source_file': filename,
                    'text_length': len(cv_text),
                    'score_extraction': score_extraction,
                    'score_matching': score_matching,
                    'score_final': score_final,
                    'extracted_info': nlp_info
                }
                
                # 🔴 CORRECTION CRITIQUE: Nettoyer tous les champs texte avant stockage
                # Mettre à jour candidature
                candidature.parsed_json = parsed_data
                candidature.score = float(score_final)
                
                # Stocker échantillon nettoyé
                cv_text_sample = cv_text[:2000]
                cv_text_sample = cv_text_sample.replace('\x00', ' ').replace('\0', ' ')
                candidature.cv_text = cv_text_sample
                
                # Mettre à jour informations extraites depuis NLP
                if nlp_info and 'structured' in nlp_info:
                    structured = nlp_info.get('structured', {})
                    
                    # Mettre à jour nom complet (nettoyé)
                    if structured.get('fullname'):
                        fullname = str(structured['fullname']).replace('\x00', ' ').replace('\0', ' ')
                        candidature.fullname = fullname[:100]  # Limiter la longueur
                        logger.info(f"   👤 Nom extrait: {fullname}")
                    
                    # Mettre à jour téléphone (nettoyé)
                    if structured.get('phone'):
                        phone = str(structured['phone']).replace('\x00', ' ').replace('\0', ' ')
                        candidature.phone = phone[:20]  # Limiter la longueur
                        logger.info(f"   📞 Téléphone extrait: {phone}")
                    
                    # Mettre à jour email (nettoyé)
                    if structured.get('email') and not candidature.email:
                        email = str(structured['email']).replace('\x00', ' ').replace('\0', ' ')
                        candidature.email = email[:100]  # Limiter la longueur
                    
                    # Mettre à jour compétences (nettoyé)
                    if structured.get('skills'):
                        all_skills = []
                        for category, skills in structured['skills'].items():
                            if isinstance(skills, list):
                                # Nettoyer chaque compétence
                                cleaned_skills = [str(s).replace('\x00', ' ').replace('\0', ' ') for s in skills]
                                all_skills.extend(cleaned_skills)
                        
                        if all_skills:
                            competences_text = ', '.join(all_skills[:10])
                            competences_text = competences_text.replace('\x00', ' ').replace('\0', ' ')
                            candidature.competences = competences_text[:500]  # Limiter à 500 caractères
                            logger.info(f"   🔧 Compétences extraites: {len(all_skills)}")
                    
                    # Mettre à jour années d'expérience
                    if structured.get('experience_years'):
                        try:
                            exp_years = int(structured['experience_years'])
                            candidature.experience_years = exp_years
                            logger.info(f"   📅 Expérience: {exp_years} ans")
                        except (ValueError, TypeError):
                            pass
                
                # 🔴 CORRECTION: Nettoyer les données NLP avant stockage
                if isinstance(nlp_info, dict):
                    # Fonction récursive pour nettoyer tous les strings dans un dictionnaire
                    def clean_dict(obj):
                        if isinstance(obj, dict):
                            return {k: clean_dict(v) for k, v in obj.items()}
                        elif isinstance(obj, list):
                            return [clean_dict(item) for item in obj]
                        elif isinstance(obj, str):
                            return obj.replace('\x00', ' ').replace('\0', ' ')
                        else:
                            return obj
                    
                    nlp_info_cleaned = clean_dict(nlp_info)
                    candidature.nlp_data = nlp_info_cleaned
                
            else:
                logger.warning(f"   ⚠️ Texte CV trop court ou impossible à extraire")
                score_final = 10
                candidature.score = 10.0
                nlp_info = {'confidence': 10, 'error': 'Texte trop court'}
                
        except Exception as e:
            logger.error(f"   ❌ Erreur traitement texte CV: {e}")
            score_final = 10
            candidature.score = 10.0
            nlp_info = {'confidence': 10, 'error': str(e)}
        
        # 🔥 ÉTAPE 7: SAUVEGARDE BASE DE DONNÉES
        try:
            db.commit()
            logger.info(f"✅ Candidature mise à jour avec succès")
            
            # 🔥 RETOUR RÉSULTATS
            return {
                "success": True,
                "minio_path": minio_path,
                "candidature_id": candidature_id,
                "filename": filename,
                "score": score_final,
                "nlp_info": nlp_info,
                "message": "CV traité avec succès"
            }
            
        except Exception as e:
            logger.error(f"❌ Erreur lors du commit DB: {e}")
            db.rollback()
            
            # 🔴 CORRECTION: Essayer une sauvegarde simplifiée si l'erreur persiste
            try:
                # Sauvegarder uniquement les informations essentielles
                candidature.score = float(score_final)
                candidature.raw_cv_s3 = minio_path
                
                # Nettoyer encore plus agressivement
                if hasattr(candidature, 'fullname') and candidature.fullname:
                    candidature.fullname = str(candidature.fullname).encode('ascii', 'ignore').decode('ascii')[:50]
                
                db.commit()
                logger.info(f"✅ Candidature sauvegardée en mode simplifié")
                
                return {
                    "success": True,
                    "minio_path": minio_path,
                    "candidature_id": candidature_id,
                    "filename": filename,
                    "score": score_final,
                    "nlp_info": {},
                    "message": "CV traité en mode simplifié"
                }
                
            except Exception as e2:
                logger.error(f"❌ Échec complet sauvegarde DB: {e2}")
                return {
                    "success": False,
                    "error": f"Erreur base de données: {str(e2)}"
                }
            
    except Exception as e:
        logger.error(f"❌ Erreur traitement CV: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return {"success": False, "error": str(e)}

def upload_files_to_minio(files: List[Tuple[str, bytes]], offre_ref: str, candidate_info: Dict[str, Any]) -> Dict[str, Any]:
    """Upload multiple de fichiers vers MinIO (pour batch processing)"""
    results = {
        "success": [],
        "failed": [],
        "total_files": len(files)
    }
    
    for filename, file_data in files:
        try:
            minio_path = minio_service.upload_cv(
                file_data=file_data,
                filename=filename,
                offre_ref=offre_ref,
                candidate_email=candidate_info.get('email', 'unknown@example.com')
            )
            
            if minio_path:
                results["success"].append({
                    "filename": filename,
                    "minio_path": minio_path,
                    "size_bytes": len(file_data)
                })
                logger.info(f"✅ Upload réussi: {filename}")
            else:
                results["failed"].append({
                    "filename": filename,
                    "error": "Échec upload MinIO"
                })
                logger.error(f"❌ Échec upload: {filename}")
                
        except Exception as e:
            results["failed"].append({
                "filename": filename,
                "error": str(e)
            })
            logger.error(f"❌ Erreur upload {filename}: {e}")
    
    results["success_count"] = len(results["success"])
    results["failed_count"] = len(results["failed"])
    
    return results

# ==================== FONCTIONS UTILITAIRES ====================

def test_minio_connection() -> Dict[str, Any]:
    """Tester la connexion MinIO"""
    if not minio_service.minio_available:
        return {
            "status": "error",
            "message": "Service MinIO non disponible",
            "endpoint": minio_service.endpoint
        }
    
    try:
        # Créer un fichier test
        test_content = b"Test de connexion MinIO - CV de test"
        test_filename = "test_cv.pdf"
        
        minio_path = minio_service.upload_cv(
            file_data=test_content,
            filename=test_filename,
            offre_ref="TEST_CONNECTION",
            candidate_email="test@example.com"
        )
        
        if minio_path:
            return {
                "status": "success",
                "message": "Connexion MinIO fonctionnelle",
                "test_file": minio_path,
                "endpoint": minio_service.endpoint,
                "bucket": minio_service.bucket_name
            }
        else:
            return {
                "status": "warning",
                "message": "MinIO accessible mais upload échoué",
                "endpoint": minio_service.endpoint
            }
            
    except Exception as e:
        return {
            "status": "error",
            "message": f"Erreur test MinIO: {str(e)}",
            "endpoint": minio_service.endpoint
        }

def get_file_info(file_data: bytes, filename: str) -> Dict[str, Any]:
    """Obtenir des informations sur un fichier"""
    return {
        "filename": filename,
        "size_bytes": len(file_data),
        "extension": Path(filename).suffix.lower(),
        "is_cv": any(filename.lower().endswith(ext) for ext in ['.pdf', '.doc', '.docx', '.odt']),
        "estimated_pages": len(file_data) // 5000 if filename.lower().endswith('.pdf') else None
    }

# Fonction pour nettoyer les chaînes de caractères
def clean_string(text: str) -> str:
    """
    Nettoyer une chaîne de caractères pour la base de données
    Supprime les caractères NULL et autres caractères problématiques
    """
    if not text:
        return text
    
    # Remplacer les caractères NULL
    text = str(text).replace('\x00', ' ').replace('\0', ' ')
    
    # Supprimer les autres caractères de contrôle (optionnel)
    import string
    printable = set(string.printable)
    text = ''.join(filter(lambda x: x in printable, text))
    
    return text.strip()

def save_upload_file(upload_file, destination: Path) -> str:
    """
    Fonction de compatibilité pour sauvegarder un fichier uploadé localement
    Utilisée par d'autres parties du système
    """
    import shutil
    import logging
    
    logger = logging.getLogger(__name__)
    
    try:
        # Créer le répertoire parent si nécessaire
        destination.parent.mkdir(parents=True, exist_ok=True)
        
        # Sauvegarder le fichier
        with open(destination, "wb") as buffer:
            shutil.copyfileobj(upload_file.file, buffer)
        
        logger.info(f"📁 Fichier sauvegardé localement: {destination}")
        return str(destination)
        
    except Exception as e:
        logger.error(f"❌ Erreur sauvegarde fichier {destination}: {e}")
        raise